"""
Admin Panel для EMET Bot
Запуск: python admin_panel.py
Доступ: http://localhost:5000
Пароль: ADMIN_PASSWORD из .env (по умолчанию: emet2026)
"""

import os
import io
import json
import db
import time
import tempfile
import threading
from datetime import datetime, date, timedelta
from functools import wraps

from flask import (
    Flask, render_template_string, request, redirect,
    url_for, session, flash, jsonify, send_file
)
from markupsafe import escape as html_escape
from dotenv import load_dotenv

load_dotenv()

ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "emet2026")
SECRET_KEY     = os.getenv("FLASK_SECRET", "emet-secret-2026")
OPENAI_KEY     = os.getenv("OPENAI_API_KEY")
GEMINI_KEY     = os.getenv("GEMINI_API_KEY")

app = Flask(__name__)


# ─── Weekly digest scheduler ──────────────────────────────────────────────────

def _digest_scheduler():
    """Фоновый поток: отправляет дайджест каждый понедельник в 9:00."""
    import time
    while True:
        now = datetime.now()
        # weekday(): 0 = понедельник
        days_until_monday = (7 - now.weekday()) % 7
        next_monday = now.replace(hour=9, minute=0, second=0, microsecond=0)
        if days_until_monday == 0 and now.hour < 9:
            pass  # сегодня понедельник, но ещё не 9:00
        elif days_until_monday == 0 and now.hour >= 9:
            days_until_monday = 7  # следующий понедельник
        next_monday += timedelta(days=days_until_monday)
        sleep_seconds = (next_monday - now).total_seconds()
        time.sleep(max(sleep_seconds, 60))
        _send_digest_now()


threading.Thread(target=_digest_scheduler, daemon=True, name="digest-scheduler").start()

app.secret_key = SECRET_KEY
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=os.getenv("HTTPS_ENABLED", "false").lower() == "true",
    PERMANENT_SESSION_LIFETIME=timedelta(hours=8),
)

# ─── Security headers ──────────────────────────────────────────────────────────

@app.after_request
def set_security_headers(response):
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    response.headers["Referrer-Policy"] = "no-referrer"
    response.headers["Content-Security-Policy"] = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' cdn.jsdelivr.net cdnjs.cloudflare.com; "
        "style-src 'self' 'unsafe-inline' cdn.jsdelivr.net cdnjs.cloudflare.com fonts.googleapis.com; "
        "img-src 'self' data:; "
        "font-src 'self' cdnjs.cloudflare.com fonts.gstatic.com; "
        "connect-src 'self' fonts.googleapis.com fonts.gstatic.com;"
    )
    return response

# ─── Brute-force protection (PostgreSQL-backed, виживає після рестарту) ────────

LOGIN_MAX_ATTEMPTS = 5
LOGIN_LOCKOUT_SEC  = 300  # 5 хвилин


def _get_client_ip() -> str:
    return request.headers.get("X-Forwarded-For", request.remote_addr or "unknown").split(",")[0].strip()


def _ensure_login_table():
    """Таблиця створюється в main.py init_db(), але адмінка може стартувати першою."""
    try:
        db.execute(
            "CREATE TABLE IF NOT EXISTS admin_login_attempts "
            "(ip TEXT PRIMARY KEY, count INTEGER DEFAULT 0, locked_until TIMESTAMP)"
        )
    except Exception:
        pass


_ensure_login_table()


def _is_locked(ip: str) -> tuple[bool, int]:
    """Повертає (locked, seconds_left). Читає з PostgreSQL."""
    try:
        row = db.query(
            "SELECT count, locked_until FROM admin_login_attempts WHERE ip=%s",
            (ip,), fetchone=True
        )
        if not row:
            return False, 0
        locked_until = row[1]
        if locked_until and locked_until > datetime.now():
            secs = int((locked_until - datetime.now()).total_seconds())
            return True, max(secs, 0)
    except Exception:
        pass
    return False, 0


def _record_failed(ip: str):
    try:
        db.execute(
            "INSERT INTO admin_login_attempts (ip, count) VALUES (%s, 1) "
            "ON CONFLICT (ip) DO UPDATE SET count = admin_login_attempts.count + 1",
            (ip,)
        )
        row = db.query("SELECT count FROM admin_login_attempts WHERE ip=%s", (ip,), fetchone=True)
        if row and row[0] >= LOGIN_MAX_ATTEMPTS:
            db.execute(
                "UPDATE admin_login_attempts SET count=0, locked_until=%s WHERE ip=%s",
                (datetime.now() + timedelta(seconds=LOGIN_LOCKOUT_SEC), ip)
            )
    except Exception:
        pass


def _reset_attempts(ip: str):
    try:
        db.execute("DELETE FROM admin_login_attempts WHERE ip=%s", (ip,))
    except Exception:
        pass

# ─── Auth ──────────────────────────────────────────────────────────────────────

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("logged_in"):
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated


# ─── DB helpers ───────────────────────────────────────────────────────────────

def db_query(sql, params=(), fetchone=False):
    return db.query_dict(sql, params, fetchone=fetchone)


def db_exec(sql, params=()):
    db.execute(sql, params)


# ─── Cost calc (same as dashboard.py) ─────────────────────────────────────────

PRICES = {
    "gpt-4o":            {"in": 2.50,  "out": 10.00},
    "gpt-4o-mini":       {"in": 0.15,  "out": 0.60},
    "gemini-2.0-flash":  {"in": 0.10,  "out": 0.40},
    "claude-sonnet-4-6": {"in": 3.00,  "out": 15.00},
}


def calc_cost(row):
    model = (row.get("model") or "").lower()
    t_in  = row.get("tokens_in")  or 0
    t_out = row.get("tokens_out") or 0
    for key, p in PRICES.items():
        if key in model:
            return (t_in * p["in"] + t_out * p["out"]) / 1_000_000
    return 0.0


def load_stats(date_from, date_to):
    rows = db_query(
        "SELECT * FROM logs WHERE date >= %s AND date <= %s ORDER BY date",
        (date_from + " 00:00:00", date_to + " 23:59:59")
    )
    return rows


# ─── Manual upload helpers ────────────────────────────────────────────────────

def extract_text_from_bytes(filename: str, content: bytes) -> str:
    """Extract plain text from PDF, DOCX, or TXT bytes."""
    ext = filename.rsplit(".", 1)[-1].lower()
    try:
        if ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            return "".join(p.extract_text() or "" for p in reader.pages)
        elif ext == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif ext in ("txt", "md", "csv"):
            return content.decode("utf-8", errors="replace")
    except Exception as e:
        return f"[Ошибка извлечения текста: {e}]"
    return ""


def index_document(filename: str, text: str, source_label: str = "manual_upload", uploaded_by: str = "admin_panel", category: str = "kb"):
    """Add document to ChromaDB indices in a background thread.
    category: 'kb' → kb indices, 'coach' → coach indices, 'both' → all four indices.
    """
    if not text or len(text.strip()) < 50:
        return False, "Текст слишком короткий или пустой"

    _INDEX_MAP = {
        "kb":    [("data/db_index_kb_openai",    "openai"), ("data/db_index_kb_google",    "google")],
        "coach": [("data/db_index_products_openai", "openai"), ("data/db_index_products_google", "google")],
        "both":  [("data/db_index_kb_openai",    "openai"), ("data/db_index_kb_google",    "google"),
                  ("data/db_index_products_openai", "openai"), ("data/db_index_products_google", "google")],
    }
    targets = _INDEX_MAP.get(category, _INDEX_MAP["kb"])

    def _do_index():
        try:
            from langchain_core.documents import Document
            from langchain_chroma import Chroma
            from langchain_openai import OpenAIEmbeddings
            from langchain_google_genai import GoogleGenerativeAIEmbeddings
            from langchain_text_splitters import RecursiveCharacterTextSplitter

            doc = Document(page_content=text, metadata={"source": filename, "url": source_label, "folder": category})
            splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=300)
            chunks = splitter.split_documents([doc])

            for persist_dir, emb_type in targets:
                if emb_type == "openai" and OPENAI_KEY:
                    emb = OpenAIEmbeddings(model="text-embedding-3-small", openai_api_key=OPENAI_KEY)
                    Chroma(persist_directory=persist_dir, embedding_function=emb).add_documents(chunks)
                elif emb_type == "google" and GEMINI_KEY:
                    emb = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001", google_api_key=GEMINI_KEY)
                    Chroma(persist_directory=persist_dir, embedding_function=emb).add_documents(chunks)

            # Record in sync_state
            db.execute(
                "INSERT INTO sync_state (file_id, file_name, modified_time, indexed_at, uploaded_by) "
                "VALUES (%s, %s, %s, %s, %s) "
                "ON CONFLICT (file_id) DO UPDATE SET "
                "file_name = EXCLUDED.file_name, modified_time = EXCLUDED.modified_time, "
                "indexed_at = EXCLUDED.indexed_at, uploaded_by = EXCLUDED.uploaded_by",
                (
                    f"manual_{filename}_{datetime.now().timestamp():.0f}",
                    filename,
                    datetime.now().isoformat(),
                    datetime.now().isoformat(),
                    uploaded_by,
                )
            )
            print(f"Indexed '{filename}': {len(chunks)} chunks")
        except Exception as e:
            print(f"Index error for '{filename}': {e}")

    threading.Thread(target=_do_index, daemon=True).start()
    return True, "OK"


# ─── Templates ────────────────────────────────────────────────────────────────

BASE_HTML = """
<!DOCTYPE html>
<html lang="uk">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>EMET Admin</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;500;600&family=Outfit:wght@400;500;600;700&family=Plus+Jakarta+Sans:wght@400;500;600&display=swap" rel="stylesheet">
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.min.js"></script>
<style>
  *,*::before,*::after{margin:0;padding:0;box-sizing:border-box}
  :root{
    --primary:#066aab;--primary-light:#2a9fd6;--primary-dark:#055a91;
    --primary-50:#e8f4fc;--primary-100:#c5e3f6;--primary-200:#8ec7ed;
    --bg:#f0f2f5;--card:#fff;
    --sidebar-bg:#fafbfc;--text-dark:#111827;--text-mid:#374151;--text-light:#6b7280;
    --text-muted:#94a3b8;--border:#e5e7eb;--radius:16px;--radius-sm:10px;--radius-xs:8px;
    --shadow-card:0 2px 12px rgba(0,0,0,0.04);--shadow-hover:0 6px 20px rgba(0,0,0,.06);
    --sidebar-w:232px;
    --red:#d63637;--red-light:#fef2f2;--red-border:#fca5a5;
    --green:#1a9a5c;--green-light:#e6f7ee;--orange:#e67e22;
    --muted:#f0f2f5;--border-light:#eef1f5;
    --font-heading:'Outfit',sans-serif;--font-body:'Plus Jakarta Sans',sans-serif;--font-mono:'JetBrains Mono',monospace;
    --transition:.18s cubic-bezier(.4,0,.2,1);
  }
  html{font-size:14px}
  body{font-family:var(--font-body);background:var(--bg);color:var(--text-dark);overflow-x:hidden;-webkit-font-smoothing:antialiased;-moz-osx-font-smoothing:grayscale;line-height:1.6}

  @keyframes fadeInUp{from{opacity:0;transform:translateY(16px)}to{opacity:1;transform:translateY(0)}}
  .fade-in{animation:fadeInUp .5s ease both}

  /* ── Layout ─────────────────────────────────────────────── */
  .layout{display:flex;min-height:100vh}

  /* ── Sidebar — light, concept_A_plus style ──────────────── */
  .sidebar{position:fixed;left:0;top:0;bottom:0;width:var(--sidebar-w);background:var(--sidebar-bg);border-right:1px solid var(--border);display:flex;flex-direction:column;z-index:100;padding:24px 12px 16px;overflow-y:auto}
  .sidebar-logo{padding:0 8px 24px;display:flex;align-items:baseline;gap:6px}
  .sidebar-logo .brand{font-family:var(--font-heading);font-weight:700;font-size:18px;color:var(--primary);letter-spacing:-.3px}
  .sidebar-logo .sub{font-family:var(--font-heading);font-weight:400;font-size:13px;color:var(--text-light)}
  .nav-section{flex:1;display:flex;flex-direction:column;gap:2px}
  .nav-sep{height:1px;background:var(--border);margin:12px 8px}
  .nav-item{display:flex;align-items:center;gap:10px;padding:9px 12px;border-radius:var(--radius-sm);font-family:var(--font-body);font-weight:500;font-size:13px;color:#64748b;cursor:pointer;transition:all .2s ease;text-decoration:none;user-select:none}
  .nav-item:hover{background:var(--bg);color:var(--text-mid)}
  .nav-item.active{background:linear-gradient(135deg,var(--primary),var(--primary-light));color:#fff;box-shadow:0 4px 12px rgba(6,106,171,.3)}
  .nav-item.active:hover{background:linear-gradient(135deg,var(--primary),var(--primary-light))}
  .nav-item svg{flex-shrink:0}
  .sidebar-bottom{padding:12px 8px 0;border-top:1px solid var(--border);margin-top:auto}
  .nav-item-logout{display:flex;align-items:center;gap:10px;padding:9px 12px;border-radius:var(--radius-sm);font-family:var(--font-body);font-weight:500;font-size:13px;color:#94a3b8;cursor:pointer;transition:all .2s ease;text-decoration:none;user-select:none;border:1px solid var(--border);justify-content:center;margin-top:8px}
  .nav-item-logout:hover{background:#fef2f2;color:var(--red);border-color:var(--red-border)}
  .nav-item-logout svg{flex-shrink:0}

  /* Mobile hamburger */
  .mobile-toggle{display:none;position:fixed;top:16px;left:16px;z-index:200;background:var(--primary);color:#fff;border:none;width:40px;height:40px;border-radius:var(--radius-xs);font-size:20px;cursor:pointer;box-shadow:0 2px 12px rgba(6,106,171,.25);align-items:center;justify-content:center}

  /* ── Main Content ──────────────────────────────────────── */
  .main{margin-left:var(--sidebar-w);flex:1;min-height:100vh}
  .topbar{position:sticky;top:0;z-index:50;background:rgba(255,255,255,.82);backdrop-filter:blur(12px);border-bottom:1px solid var(--border);padding:16px 32px;display:flex;align-items:center;justify-content:space-between}
  .topbar h1{font-family:var(--font-heading);font-weight:700;font-size:24px;color:var(--text-dark)}
  .content-area{padding:24px 32px 40px}

  .page-title{font-family:var(--font-heading);font-size:22px;font-weight:700;color:var(--text-dark);margin-bottom:24px;letter-spacing:-.3px}

  /* ── KPI Cards ──────────────────────────────────────────── */
  .kpi-row{display:grid;grid-template-columns:repeat(4,1fr);gap:16px;margin-bottom:28px}
  .kpi{background:var(--card);border-radius:var(--radius);padding:24px 26px;box-shadow:var(--shadow-card);border:1px solid var(--border-light);transition:box-shadow var(--transition),transform var(--transition)}
  .kpi:hover{box-shadow:var(--shadow-hover);transform:translateY(-2px)}
  .kpi .val{font-family:var(--font-mono);font-size:30px;font-weight:600;color:var(--primary);line-height:1.1;letter-spacing:-.5px}
  .kpi .lbl{font-family:var(--font-heading);font-size:11px;color:var(--text-muted);margin-top:8px;text-transform:uppercase;letter-spacing:.8px;font-weight:600}
  .kpi .sub{font-family:var(--font-body);font-size:12px;color:#a0aec0;margin-top:3px}

  /* ── Cards ──────────────────────────────────────────────── */
  .card{background:var(--card);border-radius:var(--radius);padding:28px;box-shadow:var(--shadow-card);border:1px solid var(--border-light);margin-bottom:20px;transition:box-shadow var(--transition)}
  .card:hover{box-shadow:var(--shadow-hover)}
  .card h2{font-family:var(--font-heading);font-size:13px;font-weight:700;color:#6b7a90;margin-bottom:20px;text-transform:uppercase;letter-spacing:.7px}

  .charts-row{display:grid;grid-template-columns:2fr 1fr 1fr;gap:16px;margin-bottom:20px}

  /* ── Tables ─────────────────────────────────────────────── */
  table{width:100%;border-collapse:collapse;font-size:13px}
  th{text-align:left;padding:10px 16px;background:transparent;color:var(--text-muted);font-family:var(--font-heading);font-weight:600;font-size:11px;text-transform:uppercase;letter-spacing:.6px;border-bottom:1px solid var(--border)}
  td{padding:13px 16px;border-bottom:1px solid var(--border-light);vertical-align:middle;color:var(--text-mid);font-family:var(--font-body)}
  tr{transition:background var(--transition)}
  tbody tr:hover td{background:linear-gradient(90deg,rgba(6,106,171,.03),transparent)}
  tbody tr:last-child td{border-bottom:none}
  td a{color:var(--primary);text-decoration:none;font-weight:500}
  td a:hover{text-decoration:underline}
  td code{font-size:12px;background:var(--muted);padding:2px 7px;border-radius:5px;font-family:var(--font-mono);color:#475569}

  /* ── Badges ─────────────────────────────────────────────── */
  .badge{display:inline-block;padding:3px 10px;border-radius:20px;font-size:11px;font-weight:600;letter-spacing:.2px;font-family:var(--font-body)}
  .badge-coach{background:var(--primary-50);color:var(--primary)}
  .badge-kb{background:var(--green-light);color:var(--green)}
  .badge-cases{background:#fef3e2;color:var(--orange)}
  .badge-operational{background:#f3e8ff;color:#7c3aed}
  .badge-openai{background:#f3e8ff;color:#7c3aed}
  .badge-google{background:#fce4ec;color:var(--red)}
  .badge-manual{background:var(--primary-50);color:var(--primary)}
  .badge-drive{background:var(--green-light);color:var(--green)}

  /* ── Buttons ────────────────────────────────────────────── */
  .btn{display:inline-flex;align-items:center;justify-content:center;gap:6px;padding:10px 22px;border-radius:var(--radius-sm);cursor:pointer;font-family:var(--font-body);font-size:13.5px;font-weight:600;border:none;text-decoration:none;transition:all var(--transition);white-space:nowrap;line-height:1.4}
  .btn:active{transform:scale(.98)}
  .btn-primary{background:linear-gradient(135deg,var(--primary),var(--primary-light));color:#fff;box-shadow:0 2px 8px rgba(6,106,171,.20)}
  .btn-primary:hover{background:linear-gradient(135deg,var(--primary-dark),var(--primary));box-shadow:0 4px 18px rgba(6,106,171,.30);transform:translateY(-1px)}
  .btn-primary:active{transform:translateY(0) scale(.98)}
  .btn-danger{background:#fff;color:var(--red);border:1.5px solid var(--red-border)}
  .btn-danger:hover{background:var(--red-light);border-color:var(--red)}
  .btn-success{background:linear-gradient(135deg,#16a34a,var(--green));color:#fff;box-shadow:0 2px 8px rgba(26,154,92,.20)}
  .btn-success:hover{box-shadow:0 4px 18px rgba(26,154,92,.30);transform:translateY(-1px)}
  .btn-outline{background:transparent;color:var(--primary);border:1.5px solid var(--primary)}
  .btn-outline:hover{background:var(--primary);color:#fff}
  .btn-sm{padding:6px 14px;font-size:12px;border-radius:var(--radius-xs)}
  .btn[disabled],.btn:disabled{opacity:.5;cursor:not-allowed;pointer-events:none}

  /* ── Forms ──────────────────────────────────────────────── */
  .form-group{margin-bottom:18px}
  .form-group label{display:block;font-family:var(--font-body);font-size:13px;font-weight:600;margin-bottom:6px;color:#526484}
  .form-control{width:100%;padding:10px 14px;border:1.5px solid var(--border);border-radius:var(--radius-sm);font-family:var(--font-body);font-size:14px;outline:none;background:var(--card);transition:border-color var(--transition),box-shadow var(--transition);color:var(--text-dark)}
  .form-control:focus{border-color:var(--primary);box-shadow:0 0 0 3px rgba(6,106,171,.08)}
  select.form-control{appearance:none;background-image:url("data:image/svg+xml,%3Csvg width='10' height='6' viewBox='0 0 10 6' xmlns='http://www.w3.org/2000/svg'%3E%3Cpath d='M1 1l4 4 4-4' stroke='%238094ae' stroke-width='1.5' fill='none' stroke-linecap='round' stroke-linejoin='round'/%3E%3C/svg%3E");background-repeat:no-repeat;background-position:right 12px center;padding-right:32px}

  /* ── Alerts ─────────────────────────────────────────────── */
  .alert{padding:14px 18px;border-radius:var(--radius-sm);margin-bottom:16px;font-family:var(--font-body);font-size:13.5px;font-weight:500;line-height:1.5;border-left:4px solid transparent}
  .alert-success{background:var(--green-light);color:#166a3e;border-left-color:var(--green)}
  .alert-danger{background:var(--red-light);color:#b42a2a;border-left-color:var(--red)}
  .alert-info{background:var(--primary-50);color:#055a91;border-left-color:var(--primary)}

  /* ── Upload zone ────────────────────────────────────────── */
  .upload-zone{border:2px dashed var(--border);border-radius:var(--radius);padding:40px;text-align:center;cursor:pointer;transition:all var(--transition);background:var(--muted)}
  .upload-zone:hover{border-color:var(--primary-200);background:var(--primary-50)}
  .upload-zone input[type="file"]{display:none}

  /* ── Misc ───────────────────────────────────────────────── */
  .sync-status{display:flex;align-items:center;gap:12px;padding:14px 18px;background:var(--muted);border-radius:var(--radius-sm);margin-bottom:14px}
  .dot-green{width:9px;height:9px;border-radius:50%;background:var(--green);box-shadow:0 0 0 3px rgba(26,154,92,.15);flex-shrink:0}
  .dot-grey{width:9px;height:9px;border-radius:50%;background:#a0aec0;flex-shrink:0}
  .progress-bar-wrap{background:var(--muted);border-radius:10px;height:7px;margin-top:4px;overflow:hidden}
  .progress-bar{height:7px;border-radius:10px;background:linear-gradient(90deg,var(--primary),var(--primary-light));transition:width .4s cubic-bezier(.4,0,.2,1)}
  .text-muted{color:var(--text-muted);font-size:12px}
  .text-green{color:var(--green);font-weight:600}
  .text-red{color:var(--red);font-weight:600}
  pre{font-family:var(--font-mono);font-size:12.5px}

  /* ── Responsive ─────────────────────────────────────────── */
  @media(max-width:1200px){:root{--sidebar-w:210px}.content-area{padding:20px 20px 36px}}
  @media(max-width:768px){
    .mobile-toggle{display:flex;align-items:center;justify-content:center}
    .sidebar{transform:translateX(-100%);transition:transform .25s cubic-bezier(.4,0,.2,1);box-shadow:none}
    .sidebar.open{transform:translateX(0);box-shadow:8px 0 30px rgba(0,0,0,.15)}
    .main{margin-left:0;padding:0}
    .content-area{padding:68px 16px 40px}
    .kpi-row{grid-template-columns:repeat(2,1fr)}
    .charts-row{grid-template-columns:1fr}
  }
  @media(max-width:480px){.kpi-row{grid-template-columns:1fr}}

  /* Scrollbar */
  .sidebar::-webkit-scrollbar{width:4px}
  .sidebar::-webkit-scrollbar-track{background:transparent}
  .sidebar::-webkit-scrollbar-thumb{background:rgba(0,0,0,.08);border-radius:4px}
</style>
</head>
<body>
<button class="mobile-toggle" onclick="document.querySelector('.sidebar').classList.toggle('open')">
  <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="3" y1="6" x2="21" y2="6"/><line x1="3" y1="12" x2="21" y2="12"/><line x1="3" y1="18" x2="21" y2="18"/></svg>
</button>
<div class="layout">
  <aside class="sidebar">
    <div class="sidebar-logo"><span class="brand">EMET</span><span class="sub">Admin Panel</span></div>
    <nav class="nav-section">
      <a href="/" class="nav-item {{ 'active' if active=='dashboard' }}">
        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="3" width="7" height="7" rx="1.5"/><rect x="14" y="3" width="7" height="7" rx="1.5"/><rect x="3" y="14" width="7" height="7" rx="1.5"/><rect x="14" y="14" width="7" height="7" rx="1.5"/></svg>
        Дашборд
      </a>
      <a href="/knowledge" class="nav-item {{ 'active' if active=='knowledge' }}">
        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M4 19.5A2.5 2.5 0 016.5 17H20"/><path d="M6.5 2H20v20H6.5A2.5 2.5 0 014 19.5v-15A2.5 2.5 0 016.5 2z"/><line x1="8" y1="7" x2="16" y2="7"/><line x1="8" y1="11" x2="13" y2="11"/></svg>
        База знань
      </a>
      <a href="/users" class="nav-item {{ 'active' if active=='users' }}">
        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><circle cx="9" cy="7" r="4"/><path d="M3 21v-2a4 4 0 014-4h4a4 4 0 014 4v2"/><circle cx="17" cy="10" r="3"/><path d="M21 21v-1.5a3 3 0 00-3-3"/></svg>
        Користувачi
      </a>
      <a href="/learning" class="nav-item {{ 'active' if active=='learning' }}">
        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M22 10v6M2 10l10-5 10 5-10 5z"/><path d="M6 12v5c0 1.66 2.69 3 6 3s6-1.34 6-3v-5"/></svg>
        Навчання
      </a>
      <a href="/access" class="nav-item {{ 'active' if active=='access' }}">
        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M12 22s8-4 8-10V5l-8-3-8 3v7c0 6 8 10 8 10z"/><polyline points="9 12 11 14 15 10"/></svg>
        Доступи
      </a>
      <div class="nav-sep"></div>
      <a href="/quality" class="nav-item {{ 'active' if active=='quality' }}">
        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><circle cx="11" cy="11" r="8"/><line x1="21" y1="21" x2="16.65" y2="16.65"/><line x1="8" y1="11" x2="14" y2="11"/><line x1="11" y1="8" x2="11" y2="14"/></svg>
        Якiсть
      </a>
      <a href="/digest" class="nav-item {{ 'active' if active=='digest' }}">
        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M4 4h16c1.1 0 2 .9 2 2v12c0 1.1-.9 2-2 2H4c-1.1 0-2-.9-2-2V6c0-1.1.9-2 2-2z"/><polyline points="22,6 12,13 2,6"/></svg>
        Дайджест
      </a>
    </nav>
    <div class="sidebar-bottom">
      <a href="/logout" class="nav-item-logout">
        <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M9 21H5a2 2 0 01-2-2V5a2 2 0 012-2h4"/><polyline points="16 17 21 12 16 7"/><line x1="21" y1="12" x2="9" y2="12"/></svg>
        Вийти
      </a>
    </div>
  </aside>
  <div class="main">
    <div class="content-area">
      {% with messages = get_flashed_messages(with_categories=true) %}
        {% for cat, msg in messages %}
          <div class="alert alert-{{ cat }}">{{ msg }}</div>
        {% endfor %}
      {% endwith %}
      {{ content | safe }}
    </div>
  </div>
</div>
</body>
</html>
"""

LOGIN_HTML = """
<!DOCTYPE html>
<html lang="uk">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>EMET Admin</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;500;600&family=Outfit:wght@400;500;600;700&family=Plus+Jakarta+Sans:wght@400;500;600&display=swap" rel="stylesheet">
<style>
  *,*::before,*::after{margin:0;padding:0;box-sizing:border-box}
  :root{--primary:#066aab;--primary-light:#2a9fd6;--font-heading:'Outfit',sans-serif;--font-body:'Plus Jakarta Sans',sans-serif}
  body{font-family:var(--font-body);background:#0c1929;display:flex;align-items:center;justify-content:center;min-height:100vh;padding:20px;-webkit-font-smoothing:antialiased;-moz-osx-font-smoothing:grayscale}
  body::before{content:'';position:fixed;top:0;left:0;right:0;bottom:0;background:radial-gradient(ellipse at 20% 50%,rgba(6,106,171,.18) 0%,transparent 50%),radial-gradient(ellipse at 80% 20%,rgba(42,159,214,.12) 0%,transparent 50%),radial-gradient(ellipse at 50% 80%,rgba(6,106,171,.08) 0%,transparent 40%);pointer-events:none}
  @keyframes fadeInUp{from{opacity:0;transform:translateY(24px)}to{opacity:1;transform:translateY(0)}}
  .box{background:#fff;border-radius:20px;padding:48px 40px;box-shadow:0 24px 80px rgba(0,0,0,.30),0 1px 3px rgba(0,0,0,.08);width:100%;max-width:400px;position:relative;animation:fadeInUp .6s ease both}
  .logo-row{display:flex;align-items:baseline;gap:8px;margin-bottom:32px}
  .logo-row .brand{font-family:var(--font-heading);font-size:20px;font-weight:700;color:var(--primary);letter-spacing:-.3px}
  .logo-row .sub{font-family:var(--font-heading);font-size:14px;font-weight:400;color:#94a3b8}
  h1{font-family:var(--font-heading);font-size:24px;font-weight:700;color:#111827;margin-bottom:6px;letter-spacing:-.4px}
  .subtitle{color:#94a3b8;font-family:var(--font-body);font-size:14px;margin-bottom:32px;line-height:1.5}
  label{display:block;font-family:var(--font-body);font-size:13px;font-weight:600;margin-bottom:7px;color:#526484}
  input[type="password"]{width:100%;padding:12px 16px;border:1.5px solid #e5e7eb;border-radius:10px;font-family:var(--font-body);font-size:15px;margin-bottom:24px;outline:none;transition:border-color .2s,box-shadow .2s;color:#111827}
  input[type="password"]:focus{border-color:var(--primary);box-shadow:0 0 0 3px rgba(6,106,171,.10)}
  button{width:100%;padding:14px;border:none;border-radius:10px;font-family:var(--font-body);font-size:15px;font-weight:700;cursor:pointer;color:#fff;background:linear-gradient(135deg,var(--primary),var(--primary-light));box-shadow:0 4px 16px rgba(6,106,171,.25);transition:all .2s cubic-bezier(.4,0,.2,1)}
  button:hover{background:linear-gradient(135deg,#055a91,var(--primary));box-shadow:0 6px 24px rgba(6,106,171,.35);transform:translateY(-1px)}
  button:active{transform:translateY(0) scale(.98)}
  .err{color:#b42a2a;font-family:var(--font-body);font-size:13px;margin-bottom:18px;padding:12px 16px;background:#fef2f2;border-radius:10px;border-left:3px solid #d63637;line-height:1.4}
  @media(max-width:480px){.box{padding:36px 28px}}
</style>
</head>
<body>
<div class="box">
  <div class="logo-row"><span class="brand">EMET</span><span class="sub">Admin Panel</span></div>
  <h1>Вхiд</h1>
  <p class="subtitle">Увiйдiть, щоб продовжити</p>
  {% if error %}<div class="err">{{ error }}</div>{% endif %}
  <form method="post">
    <label>Пароль</label>
    <input type="password" name="password" autofocus placeholder="Введiть пароль">
    <button type="submit">Увiйти</button>
  </form>
</div>
</body>
</html>
"""


def render_page(content_html, active=""):
    from flask import render_template_string
    return render_template_string(
        BASE_HTML,
        content=content_html,
        active=active
    )


# ─── Routes: Auth ──────────────────────────────────────────────────────────────

@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    ip = _get_client_ip()
    locked, secs_left = _is_locked(ip)
    if locked:
        error = f"Забагато невдалих спроб. Спробуйте через {secs_left} сек."
        return render_template_string(LOGIN_HTML, error=error)
    if request.method == "POST":
        if request.form.get("password") == ADMIN_PASSWORD:
            _reset_attempts(ip)
            session["logged_in"] = True
            session.permanent = True
            return redirect(url_for("dashboard"))
        _record_failed(ip)
        _, secs_left = _is_locked(ip)
        if secs_left:
            error = f"Забагато невдалих спроб. Заблоковано на {secs_left} сек."
        else:
            error = "Невірний пароль"
    return render_template_string(LOGIN_HTML, error=error)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


# ─── Routes: Dashboard ────────────────────────────────────────────────────────

@app.route("/")
@login_required
def dashboard():
    today      = date.today().isoformat()
    week_ago   = (date.today() - timedelta(days=7)).isoformat()
    date_from  = request.args.get("from", week_ago)
    date_to    = request.args.get("to", today)

    rows = load_stats(date_from, date_to)
    total         = len(rows)
    unique_users  = len(set(r["user_id"] for r in rows)) if rows else 0
    found_count   = sum(1 for r in rows if r.get("found_in_db") == 1)
    found_pct     = round(found_count / total * 100) if total else 0
    total_cost    = sum(calc_cost(r) for r in rows)
    tokens_in_sum = sum(r.get("tokens_in") or 0 for r in rows)
    tokens_out_sum= sum(r.get("tokens_out") or 0 for r in rows)

    # Aggregates for charts
    by_day  = {}
    by_mode = {}
    by_eng  = {}
    for r in rows:
        d = r["date"][:10]
        by_day[d]  = by_day.get(d, 0) + 1
        m = r.get("mode") or "unknown"
        by_mode[m] = by_mode.get(m, 0) + 1
        e = r.get("ai_engine") or "unknown"
        by_eng[e]  = by_eng.get(e, 0) + 1

    # Cost by model
    cost_by_model = {}
    tkn_by_model  = {}
    for r in rows:
        m = r.get("model") or "unknown"
        cost_by_model[m]   = cost_by_model.get(m, 0.0) + calc_cost(r)
        tkn_by_model[m]    = tkn_by_model.get(m, [0, 0])
        tkn_by_model[m][0] += r.get("tokens_in")  or 0
        tkn_by_model[m][1] += r.get("tokens_out") or 0

    # Top users
    ucounts = {}; unames = {}
    for r in rows:
        uid   = r["user_id"]
        uname = r.get("username") or f"id{uid}"
        ucounts[uid] = ucounts.get(uid, 0) + 1
        unames[uid]  = f"@{uname}" if not str(uname).startswith("id") else uname
    top_users = sorted(ucounts.items(), key=lambda x: x[1], reverse=True)[:10]

    last10 = rows[-10:][::-1]

    # Feedback stats за вибраний період
    fb_rows = db_query(
        "SELECT rating, mode FROM feedback WHERE created_at >= %s AND created_at <= %s",
        (date_from + " 00:00:00", date_to + " 23:59:59")
    )
    fb_total   = len(fb_rows)
    fb_positive = sum(1 for r in fb_rows if r.get("rating") == 1)
    fb_negative = sum(1 for r in fb_rows if r.get("rating") == -1)
    fb_pct      = round(fb_positive / fb_total * 100) if fb_total else 0
    fb_by_mode  = {}
    for r in fb_rows:
        m = r.get("mode") or "unknown"
        if m not in fb_by_mode:
            fb_by_mode[m] = {"up": 0, "dn": 0}
        if r.get("rating") == 1:
            fb_by_mode[m]["up"] += 1
        else:
            fb_by_mode[m]["dn"] += 1

    j = lambda x: json.dumps(x, ensure_ascii=False)

    model_rows_html = ""
    for m in sorted(cost_by_model, key=lambda x: cost_by_model[x], reverse=True):
        cnt = sum(1 for r in rows if (r.get("model") or "unknown") == m)
        model_rows_html += (
            f"<tr><td><code>{m}</code></td><td>{cnt}</td>"
            f"<td>{tkn_by_model[m][0]:,}</td><td>{tkn_by_model[m][1]:,}</td>"
            f"<td style='font-weight:700'>${cost_by_model[m]:.4f}</td></tr>"
        )

    last10_html = ""
    for r in last10:
        uname_d = ("@" + r["username"]) if r.get("username") else str(r.get("user_id", ""))
        q_text  = (r.get("question") or "")[:70]
        found   = r.get("found_in_db")
        last10_html += (
            f"<tr>"
            f"<td style='color:#888;white-space:nowrap'>{(r.get('date') or '')[:16]}</td>"
            f"<td>{uname_d}</td>"
            f"<td><span class='badge badge-{r.get('mode','')}'>{r.get('mode','')}</span></td>"
            f"<td><span class='badge badge-{(r.get('ai_engine') or '').lower()}'>{r.get('ai_engine','')}</span></td>"
            f"<td title='{q_text}'>{q_text}</td>"
            f"<td class='{'text-green' if found else 'text-red'}'>{'✓' if found else '✗'}</td>"
            f"</tr>"
        )

    content = f"""
<div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:20px">
  <h1 class="page-title">Дашборд</h1>
  <form method="get" style="display:flex;gap:10px;align-items:center">
    <input class="form-control" style="width:140px" type="date" name="from" value="{date_from}">
    <span style="color:#888">—</span>
    <input class="form-control" style="width:140px" type="date" name="to" value="{date_to}">
    <button class="btn btn-primary btn-sm" type="submit">Применить</button>
  </form>
</div>

<div class="kpi-row">
  <div class="kpi"><div class="val">{total}</div><div class="lbl">Всего запросов</div></div>
  <div class="kpi"><div class="val">{unique_users}</div><div class="lbl">Уникальных пользователей</div></div>
  <div class="kpi"><div class="val">{found_pct}%</div><div class="lbl">Найдено в базе</div>
    <div class="sub">{found_count} из {total}</div></div>
  <div class="kpi"><div class="val">${total_cost:.4f}</div><div class="lbl">Потрачено (USD)</div>
    <div class="sub">{tokens_in_sum:,} in / {tokens_out_sum:,} out</div></div>
</div>

<div class="charts-row">
  <div class="card"><h2>Запросы по дням</h2>
    <div style="height:220px"><canvas id="cDays"></canvas></div></div>
  <div class="card"><h2>Режимы</h2>
    <div style="height:220px"><canvas id="cModes"></canvas></div></div>
  <div class="card"><h2>AI Движки</h2>
    <div style="height:220px"><canvas id="cEng"></canvas></div></div>
</div>

<div class="card"><h2>Расходы по моделям</h2>
  <table>
    <thead><tr><th>Модель</th><th>Запросов</th><th>Токенов (вход)</th><th>Токенов (выход)</th><th>Стоимость</th></tr></thead>
    <tbody>{model_rows_html}
      <tr style="font-weight:800;border-top:2px solid #eee">
        <td>ИТОГО</td><td>{total}</td><td>{tokens_in_sum:,}</td>
        <td>{tokens_out_sum:,}</td><td>${total_cost:.4f}</td>
      </tr>
    </tbody>
  </table>
</div>

<div class="kpi-row" style="margin-top:0">
  <div class="kpi"><div class="val">{fb_total}</div><div class="lbl">Оцінок отримано</div></div>
  <div class="kpi"><div class="val" style="color:#4caf50">{fb_positive}</div><div class="lbl">Позитивних</div></div>
  <div class="kpi"><div class="val" style="color:#e94560">{fb_negative}</div><div class="lbl">Негативних</div></div>
  <div class="kpi"><div class="val">{fb_pct}%</div><div class="lbl">Задоволеність</div></div>
</div>

<div class="card"><h2>Оцінки по режимах</h2>
  <table>
    <thead><tr><th>Режим</th><th>+</th><th>-</th><th>%</th></tr></thead>
    <tbody>{"".join(
      f"<tr><td>{m}</td><td style='color:#4caf50'>{v['up']}</td><td style='color:#e94560'>{v['dn']}</td>"
      f"<td>{round(v['up']/(v['up']+v['dn'])*100) if v['up']+v['dn'] else 0}%</td></tr>"
      for m, v in sorted(fb_by_mode.items(), key=lambda x: -(x[1]['up']+x[1]['dn']))
    ) if fb_by_mode else "<tr><td colspan='4' style='color:#888'>Оцінок поки немає</td></tr>"}</tbody>
  </table>
</div>

<div class="card"><h2>Последние 10 запросов</h2>
  <table>
    <thead><tr><th>Время</th><th>Пользователь</th><th>Режим</th><th>Движок</th><th>Вопрос</th><th>База</th></tr></thead>
    <tbody>{last10_html}</tbody>
  </table>
</div>

<script>
const C = ['#0f3460','#e94560','#533483','#0b6e4f','#f5a623','#2196f3','#4caf50','#ff5722'];
new Chart(document.getElementById('cDays'),{{
  type:'bar', data:{{ labels:{j(list(by_day.keys()))},
    datasets:[{{data:{j(list(by_day.values()))}, backgroundColor:'#0f3460', borderRadius:4}}]}},
  options:{{ responsive:true, maintainAspectRatio:false,
    plugins:{{legend:{{display:false}}}}, scales:{{y:{{beginAtZero:true,ticks:{{precision:0}}}}}} }}
}});
new Chart(document.getElementById('cModes'),{{
  type:'doughnut', data:{{ labels:{j(list(by_mode.keys()))},
    datasets:[{{data:{j(list(by_mode.values()))}, backgroundColor:C, borderWidth:0}}]}},
  options:{{ responsive:true, maintainAspectRatio:false, plugins:{{legend:{{position:'bottom'}}}} }}
}});
new Chart(document.getElementById('cEng'),{{
  type:'doughnut', data:{{ labels:{j(list(by_eng.keys()))},
    datasets:[{{data:{j(list(by_eng.values()))}, backgroundColor:['#6a1b9a','#c62828','#1565c0'], borderWidth:0}}]}},
  options:{{ responsive:true, maintainAspectRatio:false, plugins:{{legend:{{position:'bottom'}}}} }}
}});
</script>
"""
    return render_page(content, active="dashboard")


def _build_trash_html(trash: list) -> str:
    if not trash:
        return ""
    seen = {}
    for r in trash:
        fname = r.get("file_name", "")
        if fname not in seen:
            seen[fname] = r
    unique = list(seen.values())

    rows = ""
    for r in unique:
        days = int(r.get("days_left") or 0)
        days = max(days, 0)
        color = "#2e7d32" if days > 7 else ("#e65100" if days > 2 else "#c62828")
        del_by  = r.get("deleted_by") or "—"
        del_at  = str(r.get("deleted_at") or "")[:16]
        fname   = r.get("file_name", "")
        rid     = r["id"]
        safe_fname = html_escape(fname)
        confirm_msg = json.dumps("Відновити " + fname + "?")
        rows += (
            f"<tr>"
            f"<td>{safe_fname}</td>"
            f"<td class='text-muted'>{del_by}</td>"
            f"<td class='text-muted'>{del_at}</td>"
            f"<td><span style='color:{color};font-weight:600'>{days} дн.</span></td>"
            f"<td><a href='/knowledge/restore/{rid}' class='btn btn-success btn-sm' "
            f"onclick='return confirm({confirm_msg})'>↩ Відновити</a></td>"
            f"</tr>"
        )
    return f"""
<div class="card" style="border-left:4px solid #e94560">
  <h2>Кошик — видалені файли ({len(unique)})</h2>
  <p class="text-muted" style="margin-bottom:16px;font-size:13px">
    Файли зберігаються 30 днів після видалення. Після цього — видаляються назавжди.
  </p>
  <table>
    <thead><tr><th>Файл</th><th>Видалив</th><th>Дата видалення</th><th>Залишилось</th><th></th></tr></thead>
    <tbody>{rows}</tbody>
  </table>
</div>"""


# ─── Routes: Knowledge Base ───────────────────────────────────────────────────

@app.route("/knowledge")
@login_required
def knowledge():
    # Список проиндексированных файлов
    try:
        files = db_query(
            "SELECT file_id, file_name, modified_time, indexed_at, uploaded_by FROM sync_state ORDER BY indexed_at DESC"
        )
    except Exception:
        files = []

    # Кошик — видалені файли
    try:
        trash = db_query(
            "SELECT id, file_name, deleted_by, deleted_at, restore_deadline, "
            "EXTRACT(DAY FROM restore_deadline - NOW()) as days_left "
            "FROM deleted_chunks GROUP BY id, file_name, deleted_by, deleted_at, restore_deadline "
            "ORDER BY deleted_at DESC"
        )
    except Exception:
        trash = []

    files_html = ""
    for f in files:
        fid    = f.get("file_id", "")
        source = "manual" if fid.startswith("manual_") else "drive"
        badge  = f"<span class='badge badge-{source}'>{source}</span>"
        uploader = f.get("uploaded_by") or ("Google Drive" if source == "drive" else "—")
        fn = f.get("file_name", "")
        safe_fn = html_escape(fn)
        confirm_del = json.dumps("Видалити " + fn + " з бази знань?")
        delete_btn = (
            f"<a href='/knowledge/delete/{fid}' class='btn btn-danger btn-sm' "
            f"onclick='return confirm({confirm_del})'>Видалити</a>"
        )
        files_html += (
            f"<tr><td>{safe_fn}</td>"
            f"<td>{badge}</td>"
            f"<td class='text-muted'>{uploader}</td>"
            f"<td class='text-muted'>{(f.get('indexed_at') or '')[:16]}</td>"
            f"<td>{delete_btn}</td></tr>"
        )
    if not files_html:
        files_html = "<tr><td colspan='5' style='text-align:center;color:#aaa;padding:24px'>Файлів поки немає</td></tr>"

    content = f"""
<div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:20px">
  <h1 class="page-title">База знань</h1>
</div>

<div style="display:grid;grid-template-columns:3fr 2fr;gap:20px;margin-bottom:24px">

  <div class="card">
    <h2>Синхронізація Google Drive</h2>
    <div class="sync-status">
      <div class="dot-green"></div>
      <div>
        <div style="font-size:14px;font-weight:600">Google Drive підключено</div>
        <div class="text-muted">PDF, DOCX, Google Docs/Sheets — синхронізуються автоматично</div>
      </div>
    </div>
    <p style="font-size:13px;color:var(--text-light);margin-bottom:16px">
      Бот перевіряє зміни в Drive кожні <b>60 хвилин</b>. Нові та змінені файли переіндексуються.
    </p>
    <form method="post" action="/knowledge/sync">
      <button class="btn btn-primary" type="submit">Синхронізувати</button>
    </form>
  </div>

  <div class="card">
    <h2>Завантажити файл</h2>
    <p style="font-size:13px;color:#666;margin-bottom:16px">
      Завантажте файл напряму в базу знань. Підтримуються: <b>PDF, DOCX, TXT</b>.
      Файл буде проіндексовано.
    </p>
    <form method="post" action="/knowledge/upload" enctype="multipart/form-data" id="uploadForm">
      <div class="upload-zone" onclick="document.getElementById('fileInput').click()" id="dropZone">
        <input type="file" name="file" id="fileInput" accept=".pdf,.docx,.txt,.md"
               onchange="showFileName(this)">
        <div id="dropLabel">
          <div style="font-size:15px;font-weight:600;color:#0f3460">Оберіть файл</div>
          <div class="text-muted" style="margin-top:4px">PDF, DOCX, TXT — до 20 МБ</div>
        </div>
      </div>
      <div style="margin-top:14px">
        <label style="font-size:13px;font-weight:600;color:#444;display:block;margin-bottom:6px">Категорія індексу</label>
        <select name="category" class="form-control">
          <option value="kb">База знань (регламенти)</option>
          <option value="coach">Коуч (продажі, аргументи)</option>
          <option value="both">Обидва індекси</option>
        </select>
      </div>
      <div style="margin-top:12px">
        <button class="btn btn-primary" type="submit" style="width:100%">Завантажити та проіндексувати</button>
      </div>
    </form>
  </div>
</div>

<!-- Список файлов -->
<div class="card">
  <h2>Проіндексовані файли ({len(files)})</h2>
  <table>
    <thead><tr><th>Файл</th><th>Джерело</th><th>Завантажив</th><th>Проіндексовано</th><th></th></tr></thead>
    <tbody>{files_html}</tbody>
  </table>
</div>

<!-- Кошик -->
{_build_trash_html(trash)}

<script>
function showFileName(input) {{
  if (input.files.length > 0) {{
    document.getElementById('dropLabel').innerHTML =
      '<div style="font-weight:600;color:#2e7d32">' +
      input.files[0].name + '</div><div class="text-muted">Готово до завантаження</div>';
  }}
}}
</script>
"""
    # Fill in sync interval
    try:
        import sync_manager as sm
        interval_min = sm.SYNC_INTERVAL_SEC // 60
    except Exception:
        interval_min = 60
    content = content.replace("{{ interval_min }}", str(interval_min))

    return render_page(content, active="knowledge")


@app.route("/knowledge/upload", methods=["POST"])
@login_required
def knowledge_upload():
    f = request.files.get("file")
    if not f or not f.filename:
        flash("Файл не выбран", "danger")
        return redirect(url_for("knowledge"))

    fname   = f.filename
    content = f.read()
    if len(content) > 20 * 1024 * 1024:
        flash("Файл слишком большой (максимум 20 МБ)", "danger")
        return redirect(url_for("knowledge"))

    text = extract_text_from_bytes(fname, content)
    if not text or len(text.strip()) < 50:
        flash(f"Не удалось извлечь текст из «{fname}». Проверьте формат файла.", "danger")
        return redirect(url_for("knowledge"))

    category = request.form.get("category", "kb")
    if category not in ("kb", "coach", "both"):
        category = "kb"
    ok, msg = index_document(fname, text, source_label="manual_upload",
                             uploaded_by=session.get("username", "admin_panel"),
                             category=category)
    if ok:
        cat_label = {"kb": "База знань", "coach": "Коуч", "both": "База знань + Коуч"}.get(category, category)
        flash(f"«{fname}» відправлено на індексацію ({cat_label}). Файл з'явиться в списку за кілька секунд.", "success")
    else:
        flash(f"Ошибка индексации: {msg}", "danger")
    return redirect(url_for("knowledge"))


@app.route("/knowledge/delete/<path:file_id>")
@login_required
def knowledge_delete(file_id):
    try:
        row = db_query("SELECT file_name FROM sync_state WHERE file_id=%s", (file_id,), fetchone=True)
        fname = row["file_name"] if row else file_id

        # Soft-delete: зберігаємо чанки в БД, видаляємо з ChromaDB
        deleted_chunks = _delete_from_chroma(fname, file_id, deleted_by=session.get("username", "admin_panel"))

        # Видаляємо з sync_state
        db_exec("DELETE FROM sync_state WHERE file_id=%s", (file_id,))

        flash(f"«{fname}» видалено ({deleted_chunks} чанків з індексу).", "success")
    except Exception as e:
        flash(f"Помилка видалення: {e}", "danger")
    return redirect(url_for("knowledge"))


def _delete_from_chroma(fname: str, file_id: str, deleted_by: str = "admin_panel") -> int:
    """Soft-delete: зберігає чанки в БД, видаляє з ChromaDB. Повертає кількість видалених."""
    import json
    total = 0
    indices = [
        "data/db_index_kb_openai",
        "data/db_index_kb_google",
        "data/db_index_coach_openai",
        "data/db_index_coach_google",
    ]
    try:
        from langchain_chroma import Chroma
        from langchain_openai import OpenAIEmbeddings
        from langchain_google_genai import GoogleGenerativeAIEmbeddings

        for path in indices:
            if not os.path.exists(path):
                continue
            try:
                if "openai" in path:
                    emb = OpenAIEmbeddings(model="text-embedding-3-small", openai_api_key=OPENAI_KEY)
                else:
                    emb = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001", google_api_key=GEMINI_KEY)

                vdb = Chroma(persist_directory=path, embedding_function=emb)
                result = vdb._collection.get(
                    where={"$or": [{"source": fname}, {"source": {"$contains": fname}}]},
                    include=["documents", "metadatas"]
                )
                ids = result.get("ids", [])
                if not ids:
                    continue

                # Зберігаємо чанки в БД перед видаленням
                chunks_data = {
                    "ids": ids,
                    "documents": result.get("documents", []),
                    "metadatas": result.get("metadatas", []),
                }
                db.execute(
                    "INSERT INTO deleted_chunks (file_name, file_id, index_path, chunks_json, deleted_by, restore_deadline) "
                    "VALUES (%s, %s, %s, %s, %s, NOW() + INTERVAL '30 days')",
                    (fname, file_id, path, json.dumps(chunks_data, ensure_ascii=False), deleted_by)
                )
                vdb._collection.delete(ids=ids)
                total += len(ids)
            except Exception as e:
                print(f"[delete_chroma] {path}: {e}")
    except Exception as e:
        print(f"[delete_chroma] import error: {e}")
    return total


def _restore_to_chroma(deleted_id: int) -> tuple[str, int]:
    """Відновлює чанки з БД назад в ChromaDB. Повертає (fname, кількість)."""
    import json
    row = db.query_dict("SELECT * FROM deleted_chunks WHERE id=%s", (deleted_id,), fetchone=True)
    if not row:
        return "", 0

    fname     = row["file_name"]
    file_id   = row["file_id"]
    path      = row["index_path"]
    chunks    = json.loads(row["chunks_json"])
    total     = 0

    try:
        from langchain_chroma import Chroma
        from langchain_openai import OpenAIEmbeddings
        from langchain_google_genai import GoogleGenerativeAIEmbeddings
        from langchain_core.documents import Document

        if "openai" in path:
            emb = OpenAIEmbeddings(model="text-embedding-3-small", openai_api_key=OPENAI_KEY)
        else:
            emb = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001", google_api_key=GEMINI_KEY)

        vdb = Chroma(persist_directory=path, embedding_function=emb)
        docs = [
            Document(page_content=chunks["documents"][i], metadata=chunks["metadatas"][i])
            for i in range(len(chunks["ids"]))
        ]
        vdb.add_documents(docs)
        total = len(docs)

        # Видаляємо з кошика після відновлення
        db.execute("DELETE FROM deleted_chunks WHERE id=%s", (deleted_id,))

        # Відновлюємо sync_state (якщо запису немає)
        db.execute(
            "INSERT INTO sync_state (file_id, file_name, modified_time, indexed_at, uploaded_by) "
            "VALUES (%s,%s,%s,%s,%s) ON CONFLICT(file_id) DO NOTHING",
            (file_id, fname, datetime.now().isoformat(), datetime.now().isoformat(), "restored")
        )
    except Exception as e:
        print(f"[restore_chroma] {e}")
        raise

    return fname, total


@app.route("/knowledge/restore/<int:deleted_id>")
@login_required
def knowledge_restore(deleted_id):
    try:
        fname, count = _restore_to_chroma(deleted_id)
        flash(f"«{fname}» відновлено ({count} чанків повернуто в індекс).", "success")
    except Exception as e:
        flash(f"Помилка відновлення: {e}", "danger")
    return redirect(url_for("knowledge"))


@app.route("/knowledge/sync", methods=["POST"])
@login_required
def knowledge_sync():
    def _run():
        try:
            import sync_manager
            sync_manager.run_sync()
        except Exception as e:
            print(f"Manual sync error: {e}")

    threading.Thread(target=_run, daemon=True).start()
    flash("Синхронізацію з Google Drive запущено у фоні. Оновіть сторінку через 1-2 хвилини.", "info")
    return redirect(url_for("knowledge"))


# ─── Routes: Users ─────────────────────────────────────────────────────────────

@app.route("/users")
@login_required
def users():
    try:
        user_rows = db_query(
            "SELECT u.user_id, u.username, u.first_name, u.role, u.level, "
            "u.registered_at, u.last_active, u.is_active, "
            "COUNT(l.id) as msg_count "
            "FROM users u LEFT JOIN logs l ON l.user_id = u.user_id "
            "GROUP BY u.user_id ORDER BY u.last_active DESC"
        )
    except Exception:
        user_rows = []

    # Onboarding progress per user
    onb_progress = {}
    try:
        total_onb = db_query("SELECT COUNT(*) as c FROM onboarding_items", fetchone=True)
        total_onb = (total_onb["c"] if total_onb else 0)
        rows_onb = db_query(
            "SELECT user_id, SUM(completed) as done FROM onboarding_progress GROUP BY user_id"
        )
        onb_progress = {r["user_id"]: r["done"] for r in rows_onb}
    except Exception:
        total_onb = 0

    # Test stats per user
    test_stats = {}
    try:
        rows_tests = db_query(
            "SELECT user_id, COUNT(*) as cnt, AVG(score) as avg_score, SUM(passed) as passed "
            "FROM user_progress GROUP BY user_id"
        )
        test_stats = {r["user_id"]: dict(r) for r in rows_tests}
    except Exception:
        pass

    LEVEL_ICONS = {
        "junior": "Junior", "middle": "Middle",
        "senior": "Senior", "top": "Top", "novice": "Новачок"
    }

    rows_html = ""
    for u in user_rows:
        uid  = u["user_id"]
        name = u.get("first_name") or u.get("username") or uid
        uname_d = f"@{u['username']}" if u.get("username") else str(uid)
        level   = LEVEL_ICONS.get(u.get("level", ""), u.get("level", "—"))
        active_cls = "text-green" if u.get("is_active") else "text-red"
        active_txt = "Активен" if u.get("is_active") else "Отключён"

        onb_done  = onb_progress.get(uid, 0)
        onb_pct   = round(onb_done / total_onb * 100) if total_onb else 0
        onb_bar   = (
            f"<div class='progress-bar-wrap'><div class='progress-bar' style='width:{onb_pct}%'></div></div>"
            f"<div class='text-muted' style='margin-top:2px'>{onb_done}/{total_onb}</div>"
        )

        ts = test_stats.get(uid, {})
        test_txt = f"{ts.get('cnt',0)} тем, ср. {ts.get('avg_score') or 0:.0f}%" if ts else "—"

        rows_html += (
            f"<tr>"
            f"<td><b>{name}</b><br><span class='text-muted'>{uname_d}</span></td>"
            f"<td>{level}</td>"
            f"<td class='{active_cls}'>{active_txt}</td>"
            f"<td>{u.get('msg_count', 0)}</td>"
            f"<td>{test_txt}</td>"
            f"<td style='min-width:120px'>{onb_bar}</td>"
            f"<td class='text-muted'>{(u.get('last_active') or '')[:16]}</td>"
            f"<td class='text-muted'>{(u.get('registered_at') or '')[:10]}</td>"
            f"</tr>"
        )
    if not rows_html:
        rows_html = "<tr><td colspan='8' style='text-align:center;color:#aaa;padding:24px'>Пользователей нет</td></tr>"

    content = f"""
<div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:20px">
  <h1 class="page-title">Користувачі ({len(user_rows)})</h1>
</div>
<div class="card">
  <table>
    <thead>
      <tr>
        <th>Пользователь</th><th>Уровень</th><th>Статус</th>
        <th>Запросов</th><th>Тесты</th><th>Онбординг</th>
        <th>Последняя активность</th><th>Зарегистрирован</th>
      </tr>
    </thead>
    <tbody>{rows_html}</tbody>
  </table>
</div>
"""
    return render_page(content, active="users")


# ─── Routes: Learning (LMS) ───────────────────────────────────────────────────

def _parse_course_xlsx(file_bytes: bytes):
    """Parse xlsx course file. Returns (title, description, topics_dict, topic_order) or raises ValueError."""
    import io as _io
    import re as _re
    try:
        from openpyxl import load_workbook as _lw
    except ImportError:
        raise ValueError("openpyxl не встановлено")

    wb = _lw(_io.BytesIO(file_bytes), data_only=True)

    if "Курс" not in wb.sheetnames:
        raise ValueError("Аркуш «Курс» не знайдено. Переконайтесь, що файл у правильному форматі (скачайте шаблон).")

    ws_meta = wb["Курс"]
    title       = str(ws_meta.cell(2, 2).value or "").strip()
    description = str(ws_meta.cell(3, 2).value or "").strip()
    if not title:
        raise ValueError("Назва курсу порожня (аркуш «Курс», клітинка B2).")

    # Find data sheet: prefer "Теми і тести", fallback to any other non-meta sheet
    data_sheet = None
    for candidate in ["Теми і тести", "Теми і тести"]:
        if candidate in wb.sheetnames:
            data_sheet = candidate
            break
    if data_sheet is None:
        ignore = {"Курс", "Інструкція", "Приклад"}
        others = [s for s in wb.sheetnames if s not in ignore]
        if others:
            data_sheet = others[0]
    if data_sheet is None:
        raise ValueError("Аркуш з темами не знайдено. Потрібен аркуш «Теми і тести».")

    ws = wb[data_sheet]
    topics = {}
    topic_order = []
    last_topic_num = None

    for row_i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 2):
        if all(v is None or str(v).strip() == "" for v in row):
            continue

        raw_num = str(row[0] or "").strip().rstrip(".")
        if raw_num.isdigit():
            last_topic_num = int(raw_num)
        elif last_topic_num is not None:
            pass  # continue previous topic
        else:
            continue

        topic_num   = last_topic_num
        topic_title = str(row[1] or "").strip().rstrip(",")
        topic_text  = str(row[2] or "").strip()
        q_text      = str(row[3] or "").strip()
        opt1        = str(row[4] or "").strip()
        opt2        = str(row[5] or "").strip()
        opt3        = str(row[6] or "").strip()
        opt4        = str(row[7] or "").strip()
        correct_raw = str(row[8] or "").strip()

        if topic_num not in topics:
            if not topic_title:
                continue
            topics[topic_num] = {"title": topic_title, "content": "", "questions": []}
            topic_order.append(topic_num)

        if topic_text:
            if topics[topic_num]["content"]:
                topics[topic_num]["content"] += "\n\n" + topic_text
            else:
                topics[topic_num]["content"] = topic_text

        if q_text:
            options = [o for o in [opt1, opt2, opt3, opt4] if o]
            if len(options) < 2:
                continue
            digits = _re.sub(r"[^0-9]", "", correct_raw)
            if not digits:
                continue
            correct_idx = int(digits) - 1
            if correct_idx < 0 or correct_idx >= len(options):
                continue
            topics[topic_num]["questions"].append({
                "text": q_text,
                "options": [(opt, i == correct_idx) for i, opt in enumerate(options)],
            })

    if not topics:
        raise ValueError("Теми не знайдено. Перевірте аркуш «Теми і тести».")

    return title, description, topics, topic_order


def _save_course_to_db(title, description, topics, topic_order):
    """Insert or overwrite course in DB. Returns (course_id, overwritten: bool)."""
    from datetime import datetime as _dt
    overwritten = False
    with db.get_connection() as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT id FROM courses WHERE title=%s", (title,))
            existing = cur.fetchone()
            if existing:
                old_id = existing[0]
                cur.execute(
                    "DELETE FROM answer_options WHERE question_id IN "
                    "(SELECT q.id FROM questions q JOIN topics t ON q.topic_id=t.id WHERE t.course_id=%s)",
                    (old_id,)
                )
                cur.execute("DELETE FROM questions WHERE topic_id IN (SELECT id FROM topics WHERE course_id=%s)", (old_id,))
                cur.execute("DELETE FROM topics WHERE course_id=%s", (old_id,))
                cur.execute("DELETE FROM courses WHERE id=%s", (old_id,))
                overwritten = True

            cur.execute(
                "INSERT INTO courses (title, description, created_at) VALUES (%s,%s,%s) RETURNING id",
                (title, description, _dt.now().isoformat())
            )
            course_id = cur.fetchone()[0]

            for order, tn in enumerate(topic_order, 1):
                t = topics[tn]
                cur.execute(
                    "INSERT INTO topics (course_id, order_num, title, content) VALUES (%s,%s,%s,%s) RETURNING id",
                    (course_id, order, t["title"], t["content"])
                )
                topic_id = cur.fetchone()[0]
                for q in t["questions"]:
                    cur.execute("INSERT INTO questions (topic_id, text) VALUES (%s,%s) RETURNING id", (topic_id, q["text"]))
                    q_id = cur.fetchone()[0]
                    for opt_text, is_correct in q["options"]:
                        cur.execute(
                            "INSERT INTO answer_options (question_id, text, is_correct) VALUES (%s,%s,%s)",
                            (q_id, opt_text, 1 if is_correct else 0)
                        )
    return course_id, overwritten


@app.route("/learning/upload", methods=["POST"])
@login_required
def learning_upload():
    f = request.files.get("file")
    if not f or not f.filename:
        flash("Файл не вибрано", "danger")
        return redirect(url_for("learning"))

    if not f.filename.lower().endswith(".xlsx"):
        flash("Тільки .xlsx файли підтримуються", "danger")
        return redirect(url_for("learning"))

    data = f.read()
    if len(data) > 10 * 1024 * 1024:
        flash("Файл занадто великий (макс. 10 МБ)", "danger")
        return redirect(url_for("learning"))

    try:
        title, description, topics, topic_order = _parse_course_xlsx(data)
    except ValueError as e:
        flash(f"Помилка парсингу: {e}", "danger")
        return redirect(url_for("learning"))
    except Exception as e:
        flash(f"Не вдалося прочитати файл: {e}", "danger")
        return redirect(url_for("learning"))

    try:
        course_id, overwritten = _save_course_to_db(title, description, topics, topic_order)
    except Exception as e:
        flash(f"Помилка запису в БД: {e}", "danger")
        return redirect(url_for("learning"))

    total_q = sum(len(t["questions"]) for t in topics.values())
    action = "оновлено" if overwritten else "додано"
    flash(
        f"Курс «{title}» {action}: {len(topic_order)} тем, {total_q} питань (id={course_id})",
        "success"
    )
    return redirect(url_for("learning"))


@app.route("/learning/course/delete/<int:course_id>", methods=["POST"])
@login_required
def learning_course_delete(course_id):
    row = db_query("SELECT title FROM courses WHERE id=%s", (course_id,), fetchone=True)
    if not row:
        flash("Курс не знайдено", "danger")
        return redirect(url_for("learning"))
    title = row["title"]
    try:
        db_exec(
            "DELETE FROM answer_options WHERE question_id IN "
            "(SELECT q.id FROM questions q JOIN topics t ON q.topic_id=t.id WHERE t.course_id=%s)",
            (course_id,)
        )
        db_exec("DELETE FROM questions WHERE topic_id IN (SELECT id FROM topics WHERE course_id=%s)", (course_id,))
        db_exec("DELETE FROM topics WHERE course_id=%s", (course_id,))
        db_exec("DELETE FROM courses WHERE id=%s", (course_id,))
        flash(f"Курс «{title}» видалено", "success")
    except Exception as e:
        flash(f"Помилка видалення: {e}", "danger")
    return redirect(url_for("learning"))


@app.route("/learning/template")
@login_required
def learning_template():
    """Serve course_template.xlsx for download."""
    import os as _os
    template_path = _os.path.join(_os.path.dirname(__file__), "course_template.xlsx")
    if not _os.path.exists(template_path):
        flash("Файл шаблону не знайдено на сервері", "danger")
        return redirect(url_for("learning"))
    return send_file(template_path, as_attachment=True, download_name="course_template.xlsx",
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.route("/learning/index_courses", methods=["POST"])
@login_required
def learning_index_courses():
    """Index all course topics into ChromaDB RAG (single sequential thread)."""
    category = request.form.get("category", "coach")
    if category not in ("kb", "coach", "both"):
        category = "coach"

    topics = db_query(
        "SELECT t.id, t.title, t.content, c.title as course_title "
        "FROM topics t JOIN courses c ON c.id=t.course_id "
        "WHERE t.content IS NOT NULL AND length(trim(t.content)) > 50 "
        "ORDER BY c.id, t.order_num"
    )

    if not topics:
        flash("Тем з вмістом не знайдено", "danger")
        return redirect(url_for("learning"))

    _INDEX_MAP = {
        "kb":    [("data/db_index_kb_openai",    "openai"), ("data/db_index_kb_google",    "google")],
        "coach": [("data/db_index_products_openai", "openai"), ("data/db_index_products_google", "google")],
        "both":  [("data/db_index_kb_openai",    "openai"), ("data/db_index_kb_google",    "google"),
                  ("data/db_index_products_openai", "openai"), ("data/db_index_products_google", "google")],
    }
    targets = _INDEX_MAP.get(category, _INDEX_MAP["coach"])
    topic_list = list(topics)  # snapshot before background thread

    def _index_all():
        try:
            from langchain_core.documents import Document
            from langchain_chroma import Chroma
            from langchain_openai import OpenAIEmbeddings
            from langchain_google_genai import GoogleGenerativeAIEmbeddings
            from langchain_text_splitters import RecursiveCharacterTextSplitter

            splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=300)

            # Init one instance per target — reuse across all topics
            vdbs = {}
            for persist_dir, emb_type in targets:
                if emb_type == "openai" and OPENAI_KEY:
                    emb = OpenAIEmbeddings(model="text-embedding-3-small", openai_api_key=OPENAI_KEY)
                    vdbs[(persist_dir, emb_type)] = Chroma(persist_directory=persist_dir, embedding_function=emb)
                elif emb_type == "google" and GEMINI_KEY:
                    emb = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001", google_api_key=GEMINI_KEY)
                    vdbs[(persist_dir, emb_type)] = Chroma(persist_directory=persist_dir, embedding_function=emb)

            ok_count = 0
            for t in topic_list:
                text = f"# {t['course_title']}\n## {t['title']}\n\n{t['content']}"
                filename = f"[LMS] {t['course_title']} — {t['title']}"
                doc = Document(page_content=text, metadata={
                    "source": filename, "url": "lms_course", "folder": category
                })
                chunks = splitter.split_documents([doc])
                for key, vdb in vdbs.items():
                    vdb.add_documents(chunks)
                ok_count += 1
                print(f"Indexed LMS topic {ok_count}/{len(topic_list)}: {filename[:60]}")

            print(f"LMS indexation done: {ok_count} topics indexed into {list(vdbs.keys())}")
        except Exception as e:
            print(f"LMS index error: {e}")

    threading.Thread(target=_index_all, daemon=True, name="lms-indexer").start()

    cat_label = {"kb": "База знань", "coach": "Коуч", "both": "База знань + Коуч"}.get(category, category)
    flash(
        f"Запущено індексацію {len(topic_list)} тем у розділ «{cat_label}». "
        f"Дані з'являться в RAG за 2–3 хвилини.",
        "success"
    )
    return redirect(url_for("learning"))


# ─── Routes: Access (email whitelist) ─────────────────────────────────────────

@app.route("/learning")
@login_required
def learning():
    # Все курсы
    courses = db_query("SELECT id, title, description, COALESCE(visible, true) as visible FROM courses ORDER BY id")

    # Все темы с кол-вом вопросов
    topics_all = db_query(
        "SELECT t.id, t.course_id, t.order_num, t.title, COUNT(q.id) as q_count "
        "FROM topics t LEFT JOIN questions q ON q.topic_id = t.id "
        "GROUP BY t.id ORDER BY t.course_id, t.order_num"
    )
    topics_by_course = {}
    for t in topics_all:
        topics_by_course.setdefault(t["course_id"], []).append(t)

    # Весь прогресс пользователей
    progress_rows = db_query(
        "SELECT up.user_id, up.topic_id, up.course_id, up.passed, up.score, up.attempts, up.last_date, "
        "COALESCE(u.first_name, u.username, up.user_id) as name, u.username "
        "FROM user_progress up LEFT JOIN users u ON u.user_id = up.user_id "
        "ORDER BY up.last_date DESC"
    )

    # Индексируем прогресс: (user_id, topic_id) → row
    progress_map = {(r["user_id"], r["topic_id"]): r for r in progress_rows}

    # Все уникальные пользователи с прогрессом
    users_with_progress = {}
    for r in progress_rows:
        uid = r["user_id"]
        if uid not in users_with_progress:
            users_with_progress[uid] = {"name": r["name"], "username": r["username"]}

    # --- Строим HTML ---
    # 1. KPI сводка
    total_attempts = len(progress_rows)
    total_passed   = sum(1 for r in progress_rows if r["passed"])
    unique_learners = len(users_with_progress)

    kpi_html = f"""
<div class='kpi-row' style='grid-template-columns:repeat(3,1fr)'>
  <div class='kpi'><div class='val'>{len(courses)}</div><div class='lbl'>Курсів в системі</div></div>
  <div class='kpi'><div class='val'>{unique_learners}</div><div class='lbl'>Активних учнів</div></div>
  <div class='kpi'><div class='val'>{total_passed}/{total_attempts}</div><div class='lbl'>Тем зараховано / всього спроб</div></div>
</div>"""

    # 2. Таблиця по курсах — зведення
    course_summary_rows = ""
    for c in courses:
        cid = c["id"]
        ctopics = topics_by_course.get(cid, [])
        topic_ids = [t["id"] for t in ctopics]
        learners_in_course = set(r["user_id"] for r in progress_rows if r["course_id"] == cid)
        fully_passed = sum(
            1 for uid in learners_in_course
            if all(progress_map.get((uid, tid), {}).get("passed") for tid in topic_ids)
        )
        course_scores = [r["score"] for r in progress_rows if r["course_id"] == cid and r["score"] is not None]
        avg = round(sum(course_scores) / len(course_scores)) if course_scores else 0

        ctitle = c['title']
        cdesc  = c.get('description', '')[:60]
        is_hidden = not c.get('visible', True)
        hidden_badge = " <span style='font-size:10px;background:#fff3e0;color:#e65100;padding:2px 6px;border-radius:4px'>внутрішній</span>" if is_hidden else ""
        confirm_msg = f"Видалити курс '{ctitle}' та всі результати?"
        course_summary_rows += (
            f"<tr>"
            f"<td><b>{ctitle}</b>{hidden_badge}"
            f"<br><span style='font-size:11px;color:#999'>{cdesc}</span></td>"
            f"<td style='text-align:center'>{len(ctopics)}</td>"
            f"<td style='text-align:center'>{len(learners_in_course)}</td>"
            f"<td style='text-align:center'><span style='color:#2e7d32;font-weight:700'>{fully_passed}</span></td>"
            f"<td style='text-align:center'>{avg}%</td>"
            f"<td style='text-align:center'>"
            f"<form method='post' action='/learning/course/delete/{cid}' style='display:inline' "
            f"onsubmit='return confirm(\"{confirm_msg}\")'>"
            f"<button type='submit' class='btn btn-sm btn-danger'>Видалити</button></form>"
            f"</td>"
            f"</tr>"
        )

    course_table = f"""
<div class='card'>
  <h2>Зведення по курсах</h2>
  <table>
    <tr>
      <th>Курс</th>
      <th style='text-align:center'>Тем</th>
      <th style='text-align:center'>Учнів розпочали</th>
      <th style='text-align:center'>Пройшли повністю</th>
      <th style='text-align:center'>Середній бал</th>
      <th style='text-align:center'>Дії</th>
    </tr>
    {course_summary_rows or "<tr><td colspan='6' style='text-align:center;color:#aaa;padding:20px'>Курсів поки немає — завантажте перший!</td></tr>"}
  </table>
</div>"""

    # 3. Детальна таблиця — по кожному курсу окремо
    detail_html = ""
    for c in courses:
        cid = c["id"]
        ctopics = topics_by_course.get(cid, [])
        if not ctopics:
            continue

        # Заголовки колонок: Учень | Тема1 | Тема2 | ... | Загалом
        th_topics = "".join(
            f"<th style='text-align:center;max-width:120px;white-space:normal;font-size:11px'>{t['title'][:30]}</th>"
            for t in ctopics
        )
        th_row = f"<tr><th>Учень</th>{th_topics}<th style='text-align:center'>Загалом</th></tr>"

        # Рядки по учням
        learners_in_course = set(r["user_id"] for r in progress_rows if r["course_id"] == cid)
        user_rows_html = ""

        for uid in sorted(learners_in_course):
            uinfo = users_with_progress.get(uid, {})
            name = uinfo.get("name") or uid
            uname = f"@{uinfo['username']}" if uinfo.get("username") else ""

            topic_cells = ""
            scores = []
            for t in ctopics:
                prog = progress_map.get((uid, t["id"]))
                if prog:
                    passed = prog["passed"]
                    score  = prog["score"]
                    att    = prog["attempts"]
                    date_s = (prog["last_date"] or "")[:10]
                    scores.append(score)
                    color  = "#2e7d32" if passed else "#c62828"
                    icon   = "+" if passed else "-"
                    topic_cells += (
                        f"<td style='text-align:center'>"
                        f"<span style='color:{color};font-weight:700'>{icon} {score}%</span>"
                        f"<br><span style='font-size:10px;color:#999'>x{att} · {date_s}</span>"
                        f"</td>"
                    )
                else:
                    topic_cells += "<td style='text-align:center;color:#ccc'>—</td>"

            passed_count = sum(1 for t in ctopics if progress_map.get((uid, t["id"]), {}).get("passed"))
            avg_u = round(sum(scores) / len(scores)) if scores else 0
            total_cell = (
                f"<td style='text-align:center'>"
                f"<b>{passed_count}/{len(ctopics)}</b>"
                f"<br><span style='font-size:11px;color:#555'>{avg_u}%</span>"
                f"</td>"
            )

            user_rows_html += (
                f"<tr>"
                f"<td><b>{name}</b><br><span style='font-size:11px;color:#999'>{uname}</span></td>"
                f"{topic_cells}"
                f"{total_cell}"
                f"</tr>"
            )

        if not user_rows_html:
            user_rows_html = f"<tr><td colspan='{len(ctopics)+2}' style='text-align:center;color:#aaa;padding:16px'>Ніхто ще не проходив цей курс</td></tr>"

        detail_html += f"""
<div class='card'>
  <h2>{c['title']}</h2>
  <p style='font-size:13px;color:#666;margin-bottom:16px'>{c.get('description') or ''}</p>
  <div style='overflow-x:auto'>
  <table>
    {th_row}
    {user_rows_html}
  </table>
  </div>
</div>"""

    if not detail_html:
        detail_html = "<div class='card' style='text-align:center;color:#aaa;padding:40px'>Результатів тестів поки немає</div>"

    upload_card = """
<div class='card'>
  <h2>Завантажити курс (.xlsx)</h2>
  <p style='font-size:13px;color:#666;margin-bottom:16px'>
    Завантажте файл у форматі Excel-шаблону EMET. Якщо курс з такою назвою вже існує — він буде оновлений автоматично.
    <a href='/learning/template' style='color:#0f3460;font-weight:600'>⬇ Скачати шаблон</a>
  </p>
  <form method='post' action='/learning/upload' enctype='multipart/form-data'>
    <div style='display:flex;gap:12px;align-items:center;flex-wrap:wrap'>
      <label class='upload-zone' style='padding:16px 24px;cursor:pointer;flex:1;min-width:240px'>
        <input type='file' name='file' accept='.xlsx' onchange='this.closest("form").querySelector(".fname").textContent=this.files[0]?.name||""'>
        <span style='font-size:14px;color:#555'>Файл .xlsx</span>
        <span style='display:block;margin-top:4px;font-size:13px;color:#555'>Клікніть або перетягніть .xlsx файл</span>
        <span class='fname' style='display:block;margin-top:4px;font-size:12px;color:#0f3460;font-weight:600'></span>
      </label>
      <button type='submit' class='btn btn-primary' style='white-space:nowrap'>Завантажити курс</button>
    </div>
  </form>
</div>
<div class='card'>
  <h2>Додати теми курсів до бази знань (RAG)</h2>
  <p style='font-size:13px;color:#666;margin-bottom:16px'>
    Всі теми з навчальних курсів будуть проіндексовані у векторну базу — бот зможе відповідати на запитання
    про препарати на основі матеріалів з курсів. Оберіть розділ, куди потрапить інформація.
  </p>
  <form method='post' action='/learning/index_courses'>
    <div style='display:flex;gap:12px;align-items:center;flex-wrap:wrap'>
      <select name='category' class='form-control' style='width:auto;min-width:200px'>
        <option value='coach'>Sales Коуч (рекомендовано)</option>
        <option value='kb'>База знань</option>
        <option value='both'>База знань + Коуч</option>
      </select>
      <button type='submit' class='btn btn-primary' style='white-space:nowrap'
        onclick="return confirm('Проіндексувати всі теми курсів у RAG?')">
        Індексувати всі теми курсів
      </button>
    </div>
  </form>
</div>"""

    content = f"""
<div style='margin-bottom:24px'>
  <h1 class='page-title'>Навчання — прогрес команди</h1>
</div>
{upload_card}
{kpi_html}
{course_table}
{detail_html}
"""
    return render_template_string(BASE_HTML, content=content, active="learning")


@app.route("/access")
@login_required
def access():
    try:
        rows = db_query(
            "SELECT id, email, role, full_name, activated_by_user_id, activated_at, added_at "
            "FROM allowed_emails ORDER BY added_at DESC"
        )
    except Exception:
        rows = []

    # Список пользователей для дропдауна ручной активации
    try:
        tg_users = db_query(
            "SELECT user_id, first_name, username FROM users WHERE is_active=1 ORDER BY first_name"
        )
    except Exception:
        tg_users = []

    def user_select(email_id):
        opts = "".join(
            "<option value='{uid}'>{name} (id: {uid})</option>".format(
                uid=u["user_id"],
                name=u.get("first_name") or u.get("username") or u["user_id"]
            )
            for u in tg_users
        )
        return (
            f"<form method='post' action='/access/activate/{email_id}' "
            f"style='display:flex;gap:6px;align-items:center;min-width:220px'>"
            f"<select name='user_id' class='form-control' style='padding:4px 8px;font-size:12px'>"
            f"{opts}</select>"
            f"<button class='btn btn-success btn-sm' type='submit'>Зв'язати</button>"
            f"</form>"
        )

    role_colors = {
        "admin":    "#fce4ec;color:#c62828",
        "director": "#fff8e1;color:#f57f17",
        "manager":  "#e3f2fd;color:#0d47a1",
        "operator": "#f3e5f5;color:#6a1b9a",
    }

    rows_html = ""
    for r in rows:
        activated = r.get("activated_by_user_id")
        if activated:
            status_cell = "<span class='badge' style='background:#e8f5e9;color:#2e7d32'>Активний</span>"
        else:
            status_cell = user_select(r["id"]) if tg_users else "<span class='badge' style='background:#fff3e0;color:#e65100'>Очікує</span>"

        role_style = role_colors.get(r.get("role", ""), "#f5f5f5;color:#333")
        role_badge = f"<span class='badge' style='background:{role_style}'>{r.get('role','')}</span>"
        act_date = str(r.get("activated_at") or "")[:16] or "—"
        em  = r.get('email', '')
        confirm_del_em = json.dumps("Видалити " + em + "?")
        safe_em = html_escape(em)
        rid = r['id']
        rows_html += (
            f"<tr>"
            f"<td>{safe_em}</td>"
            f"<td>{role_badge}</td>"
            f"<td>{r.get('full_name') or '—'}</td>"
            f"<td>{status_cell}</td>"
            f"<td class='text-muted'>{activated or '—'}</td>"
            f"<td class='text-muted'>{act_date}</td>"
            f"<td><a href='/access/delete/{rid}' class='btn btn-danger btn-sm' "
            f"onclick='return confirm({confirm_del_em})'>✕</a></td>"
            f"</tr>"
        )
    if not rows_html:
        rows_html = "<tr><td colspan='7' style='text-align:center;color:#aaa;padding:24px'>Список порожній. Завантажте Excel або додайте вручну.</td></tr>"

    activated_count = sum(1 for r in rows if r.get("activated_by_user_id"))
    pending_count = len(rows) - activated_count

    content = f"""
<div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:20px">
  <h1 class="page-title">Управління доступами ({len(rows)})</h1>
</div>

<div class="kpi-row" style="grid-template-columns:repeat(3,1fr)">
  <div class="kpi"><div class="val">{len(rows)}</div><div class="lbl">Всього email</div></div>
  <div class="kpi"><div class="val text-green">{activated_count}</div><div class="lbl">Активовано</div></div>
  <div class="kpi"><div class="val" style="color:#e65100">{pending_count}</div><div class="lbl">Очікують</div></div>
</div>

<div style="display:grid;grid-template-columns:1fr 1fr;gap:20px;margin-bottom:20px">

  <!-- Excel завантаження -->
  <div class="card">
    <h2>Завантажити список з Excel</h2>
    <p style="font-size:13px;color:#666;margin-bottom:16px">
      Файл Excel (.xlsx) або CSV з колонками:<br>
      <b>email</b> (обов'язково) · <b>role</b> (admin/manager/operator) · <b>name</b> (ім'я, необов'язково)<br>
      Перший рядок — заголовки. Дублікати оновлюються.
    </p>
    <form method="post" action="/access/upload" enctype="multipart/form-data">
      <div class="upload-zone" onclick="document.getElementById('excelInput').click()" style="padding:24px">
        <input type="file" name="file" id="excelInput" accept=".xlsx,.csv"
               onchange="document.getElementById('excelLabel').textContent=this.files[0].name">
        <div>
          <div style="font-size:14px;margin-bottom:6px;color:#555">Файл</div>
          <div style="font-size:14px;font-weight:600;color:#0f3460" id="excelLabel">Обрати файл .xlsx або .csv</div>
        </div>
      </div>
      <button class="btn btn-primary" type="submit" style="width:100%;margin-top:14px">Завантажити</button>
    </form>
  </div>

  <!-- Додати вручну -->
  <div class="card">
    <h2>Додати вручну</h2>
    <form method="post" action="/access/add">
      <div class="form-group">
        <label>Email *</label>
        <input class="form-control" type="email" name="email" required placeholder="name@company.ua">
      </div>
      <div class="form-group">
        <label>Роль</label>
        <select class="form-control" name="role">
          <option value="manager">manager — менеджер</option>
          <option value="director">director — директор з продажів</option>
          <option value="operator">operator — оператор (завантаження контенту)</option>
          <option value="admin">admin — адміністратор</option>
        </select>
      </div>
      <div class="form-group">
        <label>Ім'я (необов'язково)</label>
        <input class="form-control" type="text" name="full_name" placeholder="Іванов Іван">
      </div>
      <button class="btn btn-primary" type="submit" style="width:100%">Додати</button>
    </form>
  </div>
</div>

<!-- Таблиця -->
<div class="card">
  <h2>Список дозволених email</h2>
  <table>
    <thead>
      <tr>
        <th>Email</th><th>Роль</th><th>Ім'я</th><th>Статус</th>
        <th>Telegram ID</th><th>Активовано</th><th></th>
      </tr>
    </thead>
    <tbody>{rows_html}</tbody>
  </table>
</div>
"""
    return render_page(content, active="access")


@app.route("/access/add", methods=["POST"])
@login_required
def access_add():
    email = (request.form.get("email") or "").strip().lower()
    role  = request.form.get("role", "manager")
    name  = (request.form.get("full_name") or "").strip() or None
    if not email or "@" not in email:
        flash("Невірний email", "danger")
        return redirect(url_for("access"))
    if role not in ("admin", "manager", "operator", "director"):
        role = "manager"
    try:
        db_exec(
            "INSERT INTO allowed_emails (email, role, full_name) VALUES (%s, %s, %s) "
            "ON CONFLICT(email) DO UPDATE SET role=EXCLUDED.role, full_name=EXCLUDED.full_name",
            (email, role, name)
        )
        flash(f"Email {email} додано з роллю {role}", "success")
    except Exception as e:
        flash(f"Помилка: {e}", "danger")
    return redirect(url_for("access"))


@app.route("/access/activate/<int:email_id>", methods=["POST"])
@login_required
def access_activate(email_id):
    user_id = (request.form.get("user_id") or "").strip()
    if not user_id:
        flash("Оберіть користувача", "danger")
        return redirect(url_for("access"))
    try:
        row = db_query("SELECT email, role FROM allowed_emails WHERE id=%s", (email_id,), fetchone=True)
        if not row:
            flash("Email не знайдено", "danger")
            return redirect(url_for("access"))
        email, role = row["email"], row["role"]
        # Прив'язуємо email → telegram user
        db_exec(
            "UPDATE allowed_emails SET activated_by_user_id=%s, activated_at=NOW() WHERE id=%s",
            (user_id, email_id)
        )
        # Оновлюємо роль користувача
        db_exec(
            "UPDATE users SET role=%s, is_active=1 WHERE user_id=%s",
            (role, user_id)
        )
        flash(f"{email} активовано для користувача {user_id} з роллю {role}", "success")
    except Exception as e:
        flash(f"Помилка: {e}", "danger")
    return redirect(url_for("access"))


@app.route("/access/delete/<int:email_id>")
@login_required
def access_delete(email_id):
    try:
        db_exec("DELETE FROM allowed_emails WHERE id=%s", (email_id,))
        flash("Email видалено", "success")
    except Exception as e:
        flash(f"Помилка: {e}", "danger")
    return redirect(url_for("access"))


@app.route("/access/upload", methods=["POST"])
@login_required
def access_upload():
    f = request.files.get("file")
    if not f or not f.filename:
        flash("Файл не обрано", "danger")
        return redirect(url_for("access"))

    fname = f.filename.lower()
    content = f.read()
    rows_data = []

    try:
        if fname.endswith(".xlsx"):
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content))
            ws = wb.active
            headers = [str(c.value or "").strip() for c in next(ws.iter_rows(min_row=1, max_row=1))]
            for row in ws.iter_rows(min_row=2, values_only=True):
                r = dict(zip(headers, row))
                rows_data.append(r)
        elif fname.endswith(".csv"):
            import csv
            text = content.decode("utf-8-sig", errors="replace")
            reader = csv.DictReader(text.splitlines())
            for row in reader:
                rows_data.append(dict(row))
        else:
            flash("Підтримуються тільки .xlsx та .csv файли", "danger")
            return redirect(url_for("access"))
    except Exception as e:
        flash(f"Помилка читання файлу: {e}", "danger")
        return redirect(url_for("access"))

    def _find_col(row: dict, keywords: list):
        """Знаходить значення колонки за ключовими словами у назві (регістр-незалежно)."""
        for key in row:
            key_low = str(key).lower()
            if any(kw in key_low for kw in keywords):
                val = row[key]
                return str(val).strip() if val is not None else ""
        return ""

    added = 0
    skipped = 0
    errors = []
    valid_roles = {"admin", "manager", "operator", "director"}

    for i, row in enumerate(rows_data, start=2):
        # Email — шукаємо колонку з "email" або "mail" або "пошт"
        email = _find_col(row, ["email", "mail", "пошт"]).lower()
        if not email or "@" not in email:
            skipped += 1
            continue

        # Роль — шукаємо колонку "role" або "роль"
        role = _find_col(row, ["role", "роль"]).lower() or "manager"
        if role not in valid_roles:
            role = "manager"

        # Ім'я — спочатку перевіряємо "name"/"full_name"/"ім'я", потім збираємо First+Last
        name = _find_col(row, ["full_name", "fullname", "name", "ім'я", "имя", "повне"])
        if not name:
            first = _find_col(row, ["first", "ім'я", "имя", "firstname"])
            last  = _find_col(row, ["last", "прізвище", "фамилия", "lastname", "surname"])
            name  = f"{first} {last}".strip() or None
        else:
            name = name or None

        try:
            db_exec(
                "INSERT INTO allowed_emails (email, role, full_name) VALUES (%s, %s, %s) "
                "ON CONFLICT(email) DO UPDATE SET role=EXCLUDED.role, full_name=EXCLUDED.full_name",
                (email, role, name)
            )
            added += 1
        except Exception as e:
            errors.append(f"Рядок {i}: {e}")

    msg = f"Імпортовано: {added}"
    if skipped:
        msg += f" | Пропущено (порожні): {skipped}"
    if errors:
        msg += f" | Помилки: {len(errors)}"
        flash(msg, "info")
        for err in errors[:3]:
            flash(err, "danger")
    else:
        flash(msg, "success")
    return redirect(url_for("access"))


# ─── Routes: Digest ───────────────────────────────────────────────────────────

@app.route("/digest")
@login_required
def digest_page():
    content = _build_digest_html()
    return render_page(content, active="digest")


def _get_digest_recipients() -> list[int]:
    """Возвращает telegram user_id всех активных admin и director."""
    try:
        rows = db.query_dict(
            "SELECT user_id FROM users WHERE role IN ('admin', 'director') AND is_active=1"
        )
        return [int(r["user_id"]) for r in rows if r.get("user_id")]
    except Exception as e:
        print(f"Digest recipients error: {e}")
        return []


def _send_digest_now():
    """Отправляет дайджест всем admin/director. Вызывается вручную и по расписанию."""
    try:
        import asyncio
        from aiogram import Bot
        token = os.getenv("TELEGRAM_TOKEN")
        recipients = _get_digest_recipients()
        if not recipients:
            print("Digest: нет получателей (admin/director)")
            return
        text = _build_digest_telegram()
        async def _do():
            b = Bot(token=token)
            for uid in recipients:
                try:
                    await b.send_message(uid, text, parse_mode="Markdown")
                except Exception as e:
                    print(f"Digest send to {uid} failed: {e}")
            await b.session.close()
        asyncio.run(_do())
        print(f"Digest sent to {len(recipients)} recipients")
    except Exception as e:
        print(f"Digest send error: {e}")


@app.route("/digest/send", methods=["POST"])
@login_required
def digest_send():
    """Отправить дайджест в Telegram прямо сейчас."""
    threading.Thread(target=_send_digest_now, daemon=True).start()
    flash("Дайджест відправлено в Telegram адміністраторам і директорам.", "success")
    return redirect(url_for("digest_page"))


@app.route("/quality", methods=["GET"])
@login_required
def quality_page():
    """Quality monitoring page — run analysis on demand."""
    last_report = ""
    report_time = ""
    report_path = os.path.join(os.path.dirname(__file__), "data", "last_quality_report.txt")
    if os.path.exists(report_path):
        with open(report_path, "r", encoding="utf-8") as f:
            last_report = f.read()
        mtime = os.path.getmtime(report_path)
        report_time = datetime.fromtimestamp(mtime).strftime("%d.%m.%Y %H:%M")

    # Check if analysis is running (lock file)
    lock_path = os.path.join(os.path.dirname(__file__), "data", "quality_running.lock")
    is_running = os.path.exists(lock_path)
    # Auto-refresh if running
    refresh_meta = "<meta http-equiv='refresh' content='5'>" if is_running else ""
    running_badge = "<span style='color:#e65100;font-weight:600'>Аналіз виконується...</span>" if is_running else ""

    content = f"""
{refresh_meta}
<div style='margin-bottom:24px'>
  <h1 class='page-title'>Quality Monitor — моніторинг якості відповідей</h1>
</div>
<div class='card'>
  <h2>Запустити аналіз</h2>
  <p style='font-size:13px;color:#666;margin-bottom:16px'>
    Аналіз діалогів за останні 24 години: хибні відповіді, пропуски RAG, крос-сейл, суперечності.
    Звіт також автоматично відправляється адміну щодня о 08:00.
  </p>
  <form method='post' action='/quality/run'>
    <button type='submit' class='btn btn-primary' {'disabled' if is_running else ''}>
      {'Аналіз виконується...' if is_running else 'Запустити аналіз зараз'}
    </button>
  </form>
  {running_badge}
</div>
<div class='card'>
  <h2>Останній звіт {('(' + report_time + ')') if report_time else ''}</h2>
  <pre style='white-space:pre-wrap;font-size:13px;background:#f7f8fa;padding:16px;border-radius:8px;max-height:600px;overflow-y:auto'>{last_report or 'Звіт ще не створювався. Натисніть кнопку вище.'}</pre>
</div>
"""
    return render_template_string(BASE_HTML, content=content, active="quality")


@app.route("/quality/run", methods=["POST"])
@login_required
def quality_run():
    """Run quality analysis now."""
    lock_path = os.path.join(os.path.dirname(__file__), "data", "quality_running.lock")
    report_path = os.path.join(os.path.dirname(__file__), "data", "last_quality_report.txt")

    def _run():
        try:
            with open(lock_path, "w") as lf:
                lf.write("running")

            # Import and patch quality_monitor to not touch sys.stdout
            import quality_monitor as qm
            import db as _db
            from datetime import datetime as _dt, timedelta as _td

            # Run analysis directly without quality_monitor's stdout wrapper
            since = (_dt.now() - _td(hours=24)).strftime("%Y-%m-%d %H:%M:%S")
            dialogs = _db.query_dict(
                "SELECT id, date, user_id, username, mode, model, found_in_db, "
                "question, answer, tokens_in, tokens_out "
                "FROM logs WHERE date >= %s ORDER BY id", (since,)
            )

            all_findings = []
            if dialogs:
                for d in dialogs:
                    all_findings.extend(qm.analyze_dialog(d))
                all_findings.extend(qm.detect_contradictions(dialogs))
                report = qm.build_report(dialogs, all_findings)
            else:
                report = "No dialogs found in last 24h."

            with open(report_path, "w", encoding="utf-8") as f:
                f.write(report)
        except Exception as e:
            with open(report_path, "w", encoding="utf-8") as f:
                f.write(f"Error: {e}")
        finally:
            try:
                os.remove(lock_path)
            except Exception:
                pass

    threading.Thread(target=_run, daemon=True).start()
    flash("Аналіз запущено. Сторінка оновиться автоматично.", "success")
    return redirect(url_for("quality_page"))


def _build_digest_telegram() -> str:
    """Формирует текст дайджеста для Telegram (Markdown)."""
    today    = date.today()
    week_ago = (today - timedelta(days=7)).isoformat()
    today_s  = today.isoformat()

    rows = load_stats(week_ago, today_s)
    total        = len(rows)
    uniq_users   = len(set(r["user_id"] for r in rows))
    found_count  = sum(1 for r in rows if r.get("found_in_db") == 1)
    found_pct    = round(found_count / total * 100) if total else 0
    total_cost   = sum(calc_cost(r) for r in rows)

    by_mode = {}
    for r in rows:
        m = r.get("mode") or "unknown"
        by_mode[m] = by_mode.get(m, 0) + 1

    mode_lines = "\n".join(f"  • {m}: {c}" for m, c in
                           sorted(by_mode.items(), key=lambda x: x[1], reverse=True))

    # Test stats
    try:
        test_rows = db_query(
            "SELECT COUNT(*) as cnt, AVG(score) as avg FROM user_progress "
            "WHERE last_date >= %s", (week_ago,)
        )
        test_cnt = test_rows[0]["cnt"] if test_rows else 0
        test_avg = test_rows[0]["avg"] or 0 if test_rows else 0
    except Exception:
        test_cnt = test_avg = 0

    # Onboarding
    try:
        onb_rows = db_query(
            "SELECT COUNT(*) as cnt FROM onboarding_progress WHERE completed=1 AND completed_at >= %s",
            (week_ago,)
        )
        onb_cnt = onb_rows[0]["cnt"] if onb_rows else 0
    except Exception:
        onb_cnt = 0

    text = (
        f"📊 *Еженедельный дайджест EMET Bot*\n"
        f"_{week_ago} — {today_s}_\n\n"
        f"👥 Активных пользователей: *{uniq_users}*\n"
        f"💬 Всего запросов: *{total}*\n"
        f"🎯 Найдено в базе: *{found_pct}%* ({found_count}/{total})\n"
        f"💰 Потрачено (API): *${total_cost:.4f}*\n\n"
        f"📋 *По режимам:*\n{mode_lines if mode_lines else '  —'}\n\n"
        f"🎓 Тестов пройдено: *{test_cnt}*, средний балл *{test_avg:.0f}%*\n"
        f"🌱 Онбординг: *{onb_cnt}* пунктов выполнено за неделю\n\n"
        f"_Сгенерировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}_"
    )
    return text


def _build_digest_html() -> str:
    today    = date.today()
    week_ago = (today - timedelta(days=7)).isoformat()
    today_s  = today.isoformat()

    rows = load_stats(week_ago, today_s)
    total       = len(rows)
    uniq_users  = len(set(r["user_id"] for r in rows))
    found_count = sum(1 for r in rows if r.get("found_in_db") == 1)
    found_pct   = round(found_count / total * 100) if total else 0
    total_cost  = sum(calc_cost(r) for r in rows)

    by_mode = {}
    for r in rows:
        m = r.get("mode") or "unknown"
        by_mode[m] = by_mode.get(m, 0) + 1

    mode_rows = "".join(
        f"<tr><td>{m}</td><td>{c}</td><td>{round(c/total*100) if total else 0}%</td></tr>"
        for m, c in sorted(by_mode.items(), key=lambda x: x[1], reverse=True)
    )

    try:
        test_rows = db_query(
            "SELECT COUNT(*) as cnt, AVG(score) as avg FROM user_progress WHERE last_date >= %s",
            (week_ago,)
        )
        test_cnt = test_rows[0]["cnt"] if test_rows else 0
        test_avg = float(test_rows[0]["avg"] or 0) if test_rows else 0.0
    except Exception:
        test_cnt = test_avg = 0

    try:
        onb_rows = db_query(
            "SELECT COUNT(*) as cnt FROM onboarding_progress WHERE completed=1 AND completed_at >= %s",
            (week_ago,)
        )
        onb_cnt = onb_rows[0]["cnt"] if onb_rows else 0
    except Exception:
        onb_cnt = 0

    preview = _build_digest_telegram().replace("\n", "<br>").replace("*", "<b>").replace("_", "<i>")

    return f"""
<div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:20px">
  <h1 class="page-title">Щотижневий дайджест</h1>
  <form method="post" action="/digest/send">
    <button class="btn btn-primary" type="submit">Відправити в Telegram</button>
  </form>
</div>
<p style="font-size:13px;color:#888;margin-bottom:20px">
  Период: <b>{week_ago}</b> — <b>{today_s}</b>. Автоматически отправляется каждый понедельник в 9:00.
</p>

<div class="kpi-row">
  <div class="kpi"><div class="val">{uniq_users}</div><div class="lbl">Активных пользователей</div></div>
  <div class="kpi"><div class="val">{total}</div><div class="lbl">Всего запросов</div></div>
  <div class="kpi"><div class="val">{found_pct}%</div><div class="lbl">Найдено в базе</div></div>
  <div class="kpi"><div class="val">${total_cost:.4f}</div><div class="lbl">Потрачено (USD)</div></div>
</div>

<div style="display:grid;grid-template-columns:1fr 1fr;gap:20px">
  <div class="card">
    <h2>Статистика по режимам</h2>
    <table>
      <thead><tr><th>Режим</th><th>Запросов</th><th>Доля</th></tr></thead>
      <tbody>{mode_rows or "<tr><td colspan='3' style='color:#aaa'>Нет данных</td></tr>"}</tbody>
    </table>
    <div style="margin-top:16px;display:flex;gap:16px">
      <div class="kpi" style="flex:1;padding:14px">
        <div class="val" style="font-size:24px">{test_cnt}</div>
        <div class="lbl">Тестов за неделю</div>
        <div class="sub">средний балл {test_avg:.0f}%</div>
      </div>
      <div class="kpi" style="flex:1;padding:14px">
        <div class="val" style="font-size:24px">{onb_cnt}</div>
        <div class="lbl">Онбординг-пунктов</div>
        <div class="sub">выполнено за неделю</div>
      </div>
    </div>
  </div>

  <div class="card">
    <h2>Предпросмотр сообщения в Telegram</h2>
    <div style="background:#f7f8fa;border-radius:8px;padding:16px;font-size:13px;
                line-height:1.7;font-family:monospace;white-space:pre-wrap;color:#1a1a2e">
{_build_digest_telegram()}
    </div>
  </div>
</div>
"""


# ─── Entry point ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    port = int(os.getenv("ADMIN_PORT", "5000"))
    db_url = os.getenv("DATABASE_URL", "postgresql://emet:emet2026@localhost:5432/emet_bot")
    print(f"\nEMET Admin Panel")
    print(f"   URL:      http://localhost:{port}")
    print(f"   Пароль:   {ADMIN_PASSWORD}")
    print(f"   БД:       {db_url}\n")
    app.run(host="0.0.0.0", port=port, debug=False)
